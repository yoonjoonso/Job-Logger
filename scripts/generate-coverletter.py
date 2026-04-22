#!/usr/bin/env python3
"""generate-coverletter.py -- Generate a cover letter DOCX + PDF from DB content.

Normal usage (called by log_job.py or directly):
    python scripts/generate-coverletter.py \\
        --archetype gameplay \\
        --job-role "Senior Unity Engineer" \\
        --job-company "Acme Games" \\
        --role-id 44 \\
        --out output/acme-games-2026-04-12/

    Outputs:
        <candidate>-coverletter.docx  -- edit this: add personal note, tweak anything
        <candidate>-coverletter.pdf   -- auto-generated, matches the DOCX at time of run

    After editing the DOCX in Word, regenerate just the PDF:
        python scripts/generate-coverletter.py --pdf-from-docx output/.../<candidate>-coverletter.docx

Optional flags:
    --role-id N                  load role from DB to auto-detect modules (C++, leadership, etc.)
    --job-team "Core Systems"    adds "on the Core Systems team" to opening
    --modules leadership,...     force-include named modules (comma-separated)
    --company-note "I've been a fan..."  custom sentence before closing
    --format letter|a4           page size (default: letter)

Auto-detected modules (from role_requirements when --role-id is given):
    language_caveat_cpp    injected if C++ is an unmatched requirement
    leadership             injected if JD has leadership/senior signals
    multi_platform         injected if JD has platform SDK signals (PlayStation, Meta, etc.)
"""
from __future__ import annotations

import argparse
import re
import sqlite3
import sys
from pathlib import Path

ROOT    = Path(__file__).resolve().parent.parent
DB_PATH = ROOT / "data" / "job-log.db"
TEMPLATE_DIR = ROOT / "config" / "coverletter-templates"


# ── DB helpers ────────────────────────────────────────────────────────────────

def load_profile(archetype: str) -> dict:
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        row = conn.execute(
            "SELECT * FROM cover_letter_profiles WHERE archetype = ?", (archetype,)
        ).fetchone()
        if row is None:
            raise SystemExit(
                f"No cover letter profile for archetype '{archetype}'. "
                f"Run: node scripts/db/db-init.mjs coverletter"
            )
        row = dict(row)
        base = row.get("base_archetype")
        if base:
            base_row = conn.execute(
                "SELECT * FROM cover_letter_profiles WHERE archetype = ?", (base,)
            ).fetchone()
            if base_row:
                base_row = dict(base_row)
                for field in ("opening", "body_p1", "body_p2", "closing"):
                    if not row.get(field) and base_row.get(field):
                        row[field] = base_row[field]
        return row


def load_modules(keys: list[str], archetype: str) -> list[dict]:
    if not keys:
        return []
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        placeholders = ",".join("?" * len(keys))
        rows = conn.execute(
            f"SELECT * FROM cover_letter_modules "
            f"WHERE key IN ({placeholders}) AND (archetype = ? OR archetype = '*') "
            f"ORDER BY position",
            keys + [archetype],
        ).fetchall()
        seen: set[str] = set()
        result = []
        for r in rows:
            d = dict(r)
            if d["key"] not in seen:
                seen.add(d["key"])
                result.append(d)
        return result


def auto_detect_module_keys(role_id: int) -> list[str]:
    """Scan role_requirements to infer which modules to inject automatically."""
    detected: list[str] = []
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            reqs = conn.execute(
                "SELECT requirement_name, requirement_normalized, matched_normalized, kind "
                "FROM role_requirements WHERE role_id = ?",
                (role_id,),
            ).fetchall()
    except Exception:
        return detected

    all_names  = " ".join((r["requirement_name"]       or "") for r in reqs).lower()
    all_norms  = " ".join((r["requirement_normalized"]  or "") for r in reqs).lower()
    unmatched  = [r for r in reqs if r["matched_normalized"] is None]
    unmatched_text = " ".join((r["requirement_normalized"] or "") for r in unmatched).lower()

    # C++ caveat: unmatched requirement that is C++
    if re.search(r'\bc\+\+\b', unmatched_text):
        detected.append("language_caveat_cpp")

    # Leadership: JD mentions leadership signals
    leadership_signals = r'\b(tech lead|technical lead|team lead|lead engineer|mentor|mentoring|senior)\b'
    if re.search(leadership_signals, all_names + " " + all_norms):
        detected.append("leadership")

    # Multi-platform: JD mentions platform SDKs
    platform_signals = r'\b(playstation|psvr|meta quest|oculus|steam|pico|vive|xbox|certification)\b'
    if re.search(platform_signals, all_names + " " + all_norms):
        detected.append("multi_platform")

    return detected


def load_candidate_name() -> str:
    try:
        with sqlite3.connect(DB_PATH) as conn:
            row = conn.execute("SELECT full_name FROM profile LIMIT 1").fetchone()
            return row[0] if row else "Candidate"
    except Exception:
        return "Candidate"


# ── Content builder ───────────────────────────────────────────────────────────

def fill_slots(text: str, slots: dict[str, str]) -> str:
    for key, value in slots.items():
        text = text.replace(key, value or "")
    return text


def build_paragraphs(
    profile: dict,
    slots: dict[str, str],
    modules: list[dict],
    company_note: str,
) -> list[tuple[str, bool]]:
    """Return list of (text, is_placeholder) tuples."""
    paras: list[tuple[str, bool]] = []

    company = slots.get("{{COMPANY}}", "").strip()
    greeting_target = f"{company} Hiring Team" if company else "Hiring Team"
    paras.append((f"Dear {greeting_target},", False))

    paras.append((fill_slots(profile["opening"], slots), False))
    paras.append((fill_slots(profile["body_p1"], slots), False))

    if profile.get("body_p2"):
        paras.append((fill_slots(profile["body_p2"], slots), False))

    for mod in sorted(modules, key=lambda m: m["position"]):
        text = fill_slots(mod["text"], slots).strip()
        if text:
            paras.append((text, False))

    if company_note.strip():
        paras.append((company_note.strip(), False))

    paras.append((fill_slots(profile["closing"], slots), False))
    return paras


# ── DOCX renderer ─────────────────────────────────────────────────────────────

def render_docx(
    paragraphs: list[tuple[str, bool]],
    candidate_name: str,
    out_stem: Path,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        raise SystemExit("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_para(
        text: str,
        space_before: float = 0.0,
        italic: bool = False,
        color: tuple[int, int, int] | None = None,
    ) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name   = "Calibri"
        run.font.size   = Pt(11)
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        fmt = p.paragraph_format
        fmt.space_before      = Pt(space_before)
        fmt.space_after       = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for i, (text, _) in enumerate(paragraphs):
        space = 8.0 if i > 0 else 0.0
        add_para(text, space_before=space)

    # Signature block
    add_para("", space_before=16.0)
    add_para("Sincerely,")
    add_para(candidate_name)

    docx_path = out_stem.with_suffix(".docx")
    doc.save(str(docx_path))
    return docx_path


def convert_to_pdf(docx_path: Path) -> Path | None:
    try:
        from docx2pdf import convert
        pdf_path = docx_path.with_suffix(".pdf")
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception as exc:
        print(f"Warning: PDF conversion failed ({exc})", file=sys.stderr)
    return None


# ── CLI ───────────────────────────────────────────────────────────────────────

def cmd_pdf_from_docx(docx_path: Path) -> None:
    """Convert an already-edited DOCX to PDF. No content regeneration."""
    if not docx_path.exists():
        raise SystemExit(f"File not found: {docx_path}")
    pdf = convert_to_pdf(docx_path)
    if pdf:
        print(str(pdf))
    else:
        raise SystemExit("PDF conversion failed.")


def cmd_generate(args: argparse.Namespace) -> None:
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

    profile = load_profile(args.archetype)

    # Start with explicitly requested modules
    explicit_keys = [k.strip() for k in args.modules.split(",") if k.strip()]

    # Auto-detect from role requirements if role-id given
    auto_keys: list[str] = []
    if args.role_id is not None:
        auto_keys = auto_detect_module_keys(args.role_id)

    # Merge: explicit overrides, auto fills the rest; preserve order (auto first by position)
    all_keys = list(dict.fromkeys(auto_keys + explicit_keys))
    modules  = load_modules(all_keys, args.archetype)

    team_suffix = f" on the {args.job_team.strip()} team" if args.job_team.strip() else ""
    slots = {
        "{{ROLE}}":        args.job_role,
        "{{COMPANY}}":     args.job_company,
        "{{TEAM_SUFFIX}}": team_suffix,
    }

    paragraphs     = build_paragraphs(profile, slots, modules, args.company_note)
    candidate_name = load_candidate_name()
    name_prefix = candidate_name.replace(" ", "") or "Candidate"

    out_stem  = out_dir / f"{name_prefix}-coverletter"
    docx_path = render_docx(paragraphs, candidate_name, out_stem)
    pdf_path  = convert_to_pdf(docx_path)

    # Always report both
    print(f"DOCX -> {docx_path}")
    if pdf_path:
        print(f"PDF  -> {pdf_path}")
    else:
        print(f"PDF  -> (conversion failed; edit DOCX and run --pdf-from-docx)")

    if auto_keys:
        print(f"Auto-injected modules: {', '.join(auto_keys)}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a cover letter from DB content.")

    # pdf-from-docx shortcut
    parser.add_argument("--pdf-from-docx", metavar="DOCX_PATH",
                        help="Convert an already-edited DOCX to PDF. No other flags needed.")

    # generation flags
    parser.add_argument("--archetype",    default="general", help="gameplay|backend|liveops|vr|cyber|general")
    parser.add_argument("--role-id",      type=int, default=None, help="DB role num for auto-detecting modules")
    parser.add_argument("--job-role",     default="",  help="Role title")
    parser.add_argument("--job-company",  default="",  help="Company name")
    parser.add_argument("--job-team",     default="",  help="Team name for 'on the X team' suffix")
    parser.add_argument("--modules",      default="",  help="Force-include modules, comma-separated")
    parser.add_argument("--company-note", default="",  help="Custom sentence before closing")
    parser.add_argument("--format",       choices=["letter", "a4"], default="letter")
    parser.add_argument("--out",          default="",  help="Output directory")
    args = parser.parse_args()

    if args.pdf_from_docx:
        cmd_pdf_from_docx(Path(args.pdf_from_docx))
        return

    if not args.out:
        parser.error("--out is required when generating a cover letter.")

    cmd_generate(args)


if __name__ == "__main__":
    main()
