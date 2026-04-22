#!/usr/bin/env python3

import argparse
import json
import re
import sqlite3
import sys
from copy import deepcopy
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_TAB_ALIGNMENT


REPO_ROOT = Path(__file__).resolve().parents[1]
CONFIG_DIR = REPO_ROOT / "config"
OUTPUT_DIR = REPO_ROOT / "output"
DB_PATH = REPO_ROOT / "data" / "job-log.db"

PROFILE_CANDIDATES = [
    CONFIG_DIR / "profile.yml",
    CONFIG_DIR / "profile.example.yml",
]

# Map archetype short names to DB profile keys when they differ.
# For most users archetype name == profile key, so this can stay empty.
# Add entries here or in scripts/local/resume_config.py if your profile keys differ.
ARCHETYPE_ID_MAP = {}

ARCHETYPE_ALIASES = {}

try:
    import sys as _sys
    _sys.path.insert(0, str(REPO_ROOT / "scripts" / "local"))
    from resume_config import ROLE_ID_MAP, EDUCATION_ID_MAP, CERT_ID_MAP, LEGACY_PROJECTS
    _sys.path.pop(0)
except ImportError:
    ROLE_ID_MAP = {}
    EDUCATION_ID_MAP = {}
    CERT_ID_MAP = {}
    LEGACY_PROJECTS = {}

SKILL_ID_MAP = {
    "skill-csharp": "C#",
    "skill-python": "Python",
    "skill-unity": "Unity",
    "skill-unreal": "Unreal Engine",
    "skill-playfab": "PlayFab",
    "skill-azure-functions": "Azure Functions",
    "skill-cosmos-db": "Cosmos DB",
    "skill-redis": "Redis",
    "skill-sql": "SQL",
    "skill-rest-api": "REST APIs",
    "skill-nunit": "NUnit",
    "skill-meta-quest": "Meta Quest",
    "skill-htc-vive": "HTC Vive",
    "skill-psvr2": "PSVR2",
    "skill-hololens": "HoloLens",
}

# Maps project/experience codes to pinned bullet text overrides.
# Populate this in scripts/local/resume_config.py with your own project codes.
PINNED_TEXT_MAP = {}


# School name constants derived from local config (empty if not configured)
_PRIMARY_SCHOOL = next(
    (v.split("|")[0] for k, v in EDUCATION_ID_MAP.items() if "|B.S.|" in v or "|M.S.|" in v),
    ""
)
_SUPPLEMENTAL_SCHOOL = next(
    (v.split("|")[0] for k, v in EDUCATION_ID_MAP.items()
     if v.split("|")[0] != _PRIMARY_SCHOOL),
    ""
) if _PRIMARY_SCHOOL else ""
_PRIMARY_SCHOOL_ALLOWED_KEYS = {v for v in EDUCATION_ID_MAP.values() if v.split("|")[0] == _PRIMARY_SCHOOL}

INLINE_ARRAY_RE = re.compile(r"^\[(.*)\]$")
INLINE_OBJECT_RE = re.compile(r"^\{(.*)\}$")

LAYOUT_GROUPED_STANDARD = "grouped_standard"
LAYOUT_GROUPED_TOP_CERTS = "grouped_top_certs"
LAYOUT_STABLE_SECTION_ORDER = ["summary", "skills", "experience", "education", "certifications"]


@dataclass
class SummaryCandidate:
    id: str
    text: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    layout_affinity: list[str] = field(default_factory=list)
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)


@dataclass
class SkillItemCandidate:
    id: str
    display: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)


@dataclass
class SkillGroupCandidate:
    id: str
    label: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    layout_affinity: list[str] = field(default_factory=list)
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)
    items: list[SkillItemCandidate] = field(default_factory=list)


@dataclass
class BulletCandidate:
    id: str
    text: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    inclusion_mode: str
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)


@dataclass
class ProjectCandidate:
    id: str
    label: str
    stack: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)
    bullets: list[BulletCandidate] = field(default_factory=list)


@dataclass
class RoleCandidate:
    id: str
    company: str
    title: str
    dates: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)
    direct_bullets: list[BulletCandidate] = field(default_factory=list)
    projects: list[ProjectCandidate] = field(default_factory=list)


@dataclass
class EducationCandidate:
    id: str
    school: str
    detail: str
    year: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)


@dataclass
class CertificationCandidate:
    id: str
    name: str
    detail: str
    year: str
    score: int
    priority: int
    min_keep: bool
    drop_cost: int
    estimated_lines: int
    selection_reason: list[str] = field(default_factory=list)
    source_ids: list[str] = field(default_factory=list)


@dataclass
class CandidateContent:
    summary: list[SummaryCandidate] = field(default_factory=list)
    skill_groups: list[SkillGroupCandidate] = field(default_factory=list)
    roles: list[RoleCandidate] = field(default_factory=list)
    education: list[EducationCandidate] = field(default_factory=list)
    certifications: list[CertificationCandidate] = field(default_factory=list)


@dataclass
class LayoutCandidate:
    id: str
    section_order: list[str]
    section_visibility: dict[str, bool]
    section_caps: dict[str, int]
    score_inputs: dict[str, Any] = field(default_factory=dict)
    score: int = 0
    estimated_lines: int = 0
    estimated_pages: float = 0.0
    selection_reason: list[str] = field(default_factory=list)


@dataclass
class RebalanceAction:
    action: str
    target_id: str
    reason: str
    details: dict[str, Any] = field(default_factory=dict)


@dataclass
class ContentSufficiency:
    summary: str
    skills: str
    experience: str
    certifications: str


@dataclass
class FitState:
    page_pressure: str
    estimated_lines: int
    estimated_pages: float
    content_sufficiency: ContentSufficiency
    rebalance_actions: list[RebalanceAction] = field(default_factory=list)


@dataclass
class PlannedSection:
    section_id: str
    visible: bool
    order: int
    layout_reason: list[str] = field(default_factory=list)
    estimated_lines: int = 0
    item_ids: list[str] = field(default_factory=list)


@dataclass
class DroppedItem:
    id: str
    kind: str
    reason: str
    score: int
    details: dict[str, Any] = field(default_factory=dict)


@dataclass
class PlanValidation:
    passed: bool
    violations: list[str] = field(default_factory=list)


@dataclass
class ResumePlan:
    version: str
    archetype: str
    profile_key: str
    format: str
    max_pages: int
    inputs: dict[str, Any]
    candidate_content: CandidateContent
    candidate_layouts: list[LayoutCandidate]
    fit_state: FitState
    final_layout: str
    final_sections: list[PlannedSection]
    dropped_items: list[DroppedItem]
    validation: PlanValidation


def fail(message):
    raise SystemExit(message)


def load_yaml(path):
    raw = path.read_text(encoding="utf-8")
    try:
        import yaml  # type: ignore

        return yaml.safe_load(raw)
    except Exception:
        return parse_minimal_yaml(raw)


def parse_minimal_yaml(raw):
    lines = raw.replace("\r\n", "\n").split("\n")
    index = 0

    def strip_comment(line):
        in_single = False
        in_double = False
        escaped = False
        for i, ch in enumerate(line):
            if ch == "\\" and not escaped:
                escaped = True
                continue
            if ch == "'" and not in_double and not escaped:
                in_single = not in_single
            elif ch == '"' and not in_single and not escaped:
                in_double = not in_double
            elif ch == "#" and not in_single and not in_double:
                return line[:i].rstrip()
            escaped = False
        return line.rstrip()

    def indent_of(line):
        return len(line) - len(line.lstrip(" "))

    def skip_noise():
        nonlocal index
        while index < len(lines):
            if strip_comment(lines[index]).strip():
                return
            index += 1

    def parse_scalar(value):
        value = value.strip()
        if not value:
            return ""
        if value in ("true", "True"):
            return True
        if value in ("false", "False"):
            return False
        if value in ("null", "Null", "~"):
            return None
        if value.startswith(("'", '"')) and value.endswith(("'", '"')) and len(value) >= 2:
            return value[1:-1]
        if INLINE_ARRAY_RE.match(value):
            return parse_inline_array(value[1:-1])
        if INLINE_OBJECT_RE.match(value):
            return parse_inline_object(value[1:-1])
        if re.fullmatch(r"-?\d+", value):
            return int(value)
        return value

    def split_inline(value):
        parts = []
        current = []
        depth = 0
        in_single = False
        in_double = False
        for ch in value:
            if ch == "'" and not in_double:
                in_single = not in_single
            elif ch == '"' and not in_single:
                in_double = not in_double
            elif not in_single and not in_double:
                if ch in "[{":
                    depth += 1
                elif ch in "]}":
                    depth -= 1
                elif ch == "," and depth == 0:
                    parts.append("".join(current).strip())
                    current = []
                    continue
            current.append(ch)
        if current:
            parts.append("".join(current).strip())
        return [part for part in parts if part]

    def parse_inline_array(value):
        return [parse_scalar(part) for part in split_inline(value)]

    def parse_inline_object(value):
        result = {}
        for part in split_inline(value):
            key, sep, rest = part.partition(":")
            if not sep:
                fail(f"Invalid inline YAML object entry: {part}")
            result[key.strip()] = parse_scalar(rest.strip())
        return result

    def parse_block_scalar(base_indent, folded):
        nonlocal index
        parts = []
        while index < len(lines):
            line = lines[index]
            if not strip_comment(line).strip():
                parts.append("")
                index += 1
                continue
            indent = indent_of(line)
            if indent < base_indent:
                break
            parts.append(line[base_indent:])
            index += 1
        if not folded:
            return "\n".join(parts).rstrip()
        merged = []
        pending_blank = False
        for part in parts:
            if part == "":
                pending_blank = True
                continue
            if pending_blank and merged:
                merged.append("\n")
            elif merged and merged[-1] != "\n":
                merged.append(" ")
            merged.append(part.strip())
            pending_blank = False
        return "".join(merged).strip()

    def find_key_separator(content):
        in_single = False
        in_double = False
        depth = 0
        for i, ch in enumerate(content):
            if ch == "'" and not in_double:
                in_single = not in_single
            elif ch == '"' and not in_single:
                in_double = not in_double
            elif not in_single and not in_double:
                if ch in "[{":
                    depth += 1
                elif ch in "]}":
                    depth -= 1
                elif ch == ":" and depth == 0:
                    return i
        return -1

    def parse_block(expected_indent):
        skip_noise()
        if index >= len(lines):
            return {}
        probe = strip_comment(lines[index])
        if indent_of(probe) < expected_indent:
            return {}
        if probe.lstrip().startswith("- "):
            return parse_sequence(expected_indent)
        return parse_mapping(expected_indent)

    def parse_sequence(expected_indent):
        nonlocal index
        items = []
        while index < len(lines):
            line = strip_comment(lines[index])
            if not line.strip():
                index += 1
                continue
            indent = indent_of(line)
            if indent < expected_indent:
                break
            if indent != expected_indent or not line[indent:].startswith("- "):
                fail(f"Invalid YAML list item near line {index + 1}")
            content = line[indent + 2 :].strip()
            index += 1
            if not content:
                items.append(parse_block(expected_indent + 2))
                continue
            if find_key_separator(content) != -1:
                key, _, rest = content.partition(":")
                entry = {}
                if rest.strip() in ("|", ">"):
                    entry[key.strip()] = parse_block_scalar(expected_indent + 2, rest.strip() == ">")
                elif rest.strip():
                    entry[key.strip()] = parse_scalar(rest.strip())
                else:
                    entry[key.strip()] = parse_block(expected_indent + 2)
                child = parse_block(expected_indent + 2)
                if isinstance(child, dict):
                    entry.update(child)
                items.append(entry)
                continue
            items.append(parse_scalar(content))
        return items

    def parse_mapping(expected_indent):
        nonlocal index
        result = {}
        while index < len(lines):
            line = strip_comment(lines[index])
            if not line.strip():
                index += 1
                continue
            indent = indent_of(line)
            if indent < expected_indent:
                break
            if indent != expected_indent:
                fail(f"Invalid YAML indentation near line {index + 1}")
            content = line[indent:]
            if content.startswith("- "):
                break
            separator = find_key_separator(content)
            if separator == -1:
                fail(f'Expected "key: value" near line {index + 1}')
            key = content[:separator].strip()
            rest = content[separator + 1 :].strip()
            index += 1
            if rest == "|":
                result[key] = parse_block_scalar(expected_indent + 2, False)
            elif rest == ">":
                result[key] = parse_block_scalar(expected_indent + 2, True)
            elif rest:
                result[key] = parse_scalar(rest)
            else:
                result[key] = parse_block(expected_indent + 2)
        return result

    parsed = parse_block(0)
    skip_noise()
    return parsed


def normalize_text(value):
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def month_label(value):
    if not value:
        return "Present"
    match = re.fullmatch(r"(\d{4})-(\d{2})", str(value).strip())
    if not match:
        return str(value)
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    return f"{months[int(match.group(2)) - 1]} {match.group(1)}"


def date_range(start_date, end_date):
    return f"{month_label(start_date)} - {month_label(end_date)}"


def profile_path():
    for candidate in PROFILE_CANDIDATES:
        if candidate.exists():
            return candidate
    fail("Missing config/profile.yml — run python setup.py to create it")


def dedupe_rows(rows, key_fn):
    seen = set()
    output = []
    for row in rows:
        key = key_fn(row)
        if key in seen:
            continue
        seen.add(key)
        output.append(dict(row))
    return output


def row_value(row, *keys, default=None):
    for key in keys:
        if key is None:
            continue
        if isinstance(row, dict) and key in row and row[key] is not None:
            return row[key]
        try:
            value = row[key]
        except Exception:
            continue
        if value is not None:
            return value
    return default


def parse_bool(value, default=True):
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = normalize_text(value)
    if text in {"1", "true", "yes", "y", "approved"}:
        return True
    if text in {"0", "false", "no", "n", "hidden"}:
        return False
    return default


def parse_int(value, default=0):
    try:
        if value is None or value == "":
            return default
        return int(value)
    except Exception:
        return default


def load_db():
    if not DB_PATH.exists():
        fail(f"Missing data source: {DB_PATH}")
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    try:
        available_tables = {
            row["name"]
            for row in connection.execute("SELECT name FROM sqlite_master WHERE type IN ('table', 'view')").fetchall()
        }
        profile_rows = connection.execute("SELECT key, value FROM profile").fetchall()
        experience_rows = connection.execute(
            "SELECT id, company, title, start_date, end_date, location FROM experience ORDER BY start_date DESC, id ASC"
        ).fetchall()
        bullet_rows = connection.execute(
            """
            SELECT b.id, b.exp_id, b.text, b.tags, b.strength, e.company, e.title, e.start_date, e.end_date
            FROM bullets b
            JOIN experience e ON e.id = b.exp_id
            ORDER BY e.start_date DESC, b.id ASC
            """
        ).fetchall()
        skill_rows = connection.execute(
            """
            SELECT id, skill_name, skill_normalized, category, display_category, secondary_categories, level, evidence,
                   resume_priority, include_default, require_direct_match, profile_bias,
                   resume_display, resume_visibility
            FROM skills_mine
            ORDER BY category, skill_name
            """
        ).fetchall()
        education_rows = connection.execute(
            "SELECT id, school, degree, field, status, grad_year FROM education ORDER BY id ASC"
        ).fetchall()
        certification_rows = connection.execute(
            "SELECT id, name, issuer, date, status FROM certifications ORDER BY date DESC, id ASC"
        ).fetchall()
        archetype_rows = connection.execute("SELECT id, name FROM archetypes ORDER BY id ASC").fetchall() if "archetypes" in available_tables else []
        project_rows = connection.execute("SELECT * FROM projects").fetchall() if "projects" in available_tables else []
        resume_point_rows = connection.execute("SELECT * FROM resume_points").fetchall() if "resume_points" in available_tables else []
        resume_point_variant_rows = (
            connection.execute("SELECT * FROM resume_point_variants").fetchall()
            if "resume_point_variants" in available_tables
            else []
        )
        resume_point_archetype_rows = (
            connection.execute("SELECT * FROM resume_point_archetypes").fetchall()
            if "resume_point_archetypes" in available_tables
            else []
        )
        resume_point_skill_rows = (
            connection.execute("SELECT * FROM resume_point_skills").fetchall()
            if "resume_point_skills" in available_tables
            else []
        )
        project_skill_rows = (
            connection.execute(
                """
                SELECT ps.*, sm.skill_name
                FROM project_skills ps
                LEFT JOIN skills_mine sm ON sm.skill_normalized = ps.skill_normalized
                """
            ).fetchall()
            if "project_skills" in available_tables
            else []
        )
        skill_alias_rows = (
            connection.execute("SELECT skill_normalized, alias_name, alias_normalized, notes FROM skill_aliases").fetchall()
            if "skill_aliases" in available_tables
            else []
        )
        skill_resume_rule_rows = (
            connection.execute("SELECT * FROM skill_resume_rules").fetchall()
            if "skill_resume_rules" in available_tables
            else []
        )
        resume_group_rule_rows = (
            connection.execute("SELECT * FROM resume_group_rules").fetchall()
            if "resume_group_rules" in available_tables
            else []
        )
        resume_layout_rows = (
            connection.execute("SELECT * FROM resume_layouts").fetchall()
            if "resume_layouts" in available_tables
            else []
        )
        resume_layout_profile_rule_rows = (
            connection.execute("SELECT * FROM resume_layout_profile_rules").fetchall()
            if "resume_layout_profile_rules" in available_tables
            else []
        )
        resume_layout_signal_rule_rows = (
            connection.execute("SELECT * FROM resume_layout_signal_rules").fetchall()
            if "resume_layout_signal_rules" in available_tables
            else []
        )
        resume_profile_rows = connection.execute("SELECT * FROM resume_profiles").fetchall() if "resume_profiles" in available_tables else []
        resume_profile_signal_rule_rows = (
            connection.execute("SELECT * FROM resume_profile_signal_rules").fetchall()
            if "resume_profile_signal_rules" in available_tables
            else []
        )
        role_rows = (
            connection.execute(
                """
                SELECT num, company, role, notes, location_text, work_model, compensation_text
                FROM roles
                """
            ).fetchall()
            if "roles" in available_tables
            else []
        )
        role_requirement_rows = (
            connection.execute(
                """
                SELECT id, role_id, raw_text, requirement_name, requirement_normalized, kind, priority,
                       matched_entity_type, matched_normalized, match_method, confidence, source, notes
                FROM role_requirements
                """
            ).fetchall()
            if "role_requirements" in available_tables
            else []
        )
        role_archetype_score_rows = (
            connection.execute(
                """
                SELECT role_id, archetype_key, score, rank, evidence_json,
                       requirement_score, keyword_score, title_score, approved
                FROM role_archetype_scores
                """
            ).fetchall()
            if "role_archetype_scores" in available_tables
            else []
        )
        archetype_resume_config_rows = (
            connection.execute(
                "SELECT id, archetype_key, label, template_key, subtitle, summary, approved, caution_rules, experience, skills, education, certifications FROM archetype_resume_configs WHERE approved = 1"
            ).fetchall()
            if "archetype_resume_configs" in available_tables
            else []
        )
        resume_generation_setting_rows = (
            connection.execute("SELECT key, value FROM resume_generation_settings").fetchall()
            if "resume_generation_settings" in available_tables
            else []
        )
    finally:
        connection.close()

    return {
        "available_tables": available_tables,
        "profile": {row["key"]: row["value"] for row in profile_rows},
        "experience": dedupe_rows(experience_rows, lambda row: (row["company"], row["title"], row["start_date"], row["end_date"])),
        "bullets": dedupe_rows(bullet_rows, lambda row: (row["exp_id"], row["text"])),
        "skills": dedupe_rows(skill_rows, lambda row: row["skill_name"]),
        "education": dedupe_rows(education_rows, lambda row: (row["school"], row["degree"], row["field"])),
        "certifications": dedupe_rows(certification_rows, lambda row: (row["name"], row["issuer"], row["date"], row["status"])),
        "archetypes": [dict(row) for row in archetype_rows],
        "projects": [dict(row) for row in project_rows],
        "resume_points": [dict(row) for row in resume_point_rows],
        "resume_point_variants": [dict(row) for row in resume_point_variant_rows],
        "resume_point_archetypes": [dict(row) for row in resume_point_archetype_rows],
        "resume_point_skills": [dict(row) for row in resume_point_skill_rows],
        "project_skills": [dict(row) for row in project_skill_rows],
        "skill_aliases": [dict(row) for row in skill_alias_rows],
        "skill_resume_rules": [dict(row) for row in skill_resume_rule_rows],
        "resume_group_rules": [dict(row) for row in resume_group_rule_rows],
        "resume_layouts": [dict(row) for row in resume_layout_rows],
        "resume_layout_profile_rules": [dict(row) for row in resume_layout_profile_rule_rows],
        "resume_layout_signal_rules": [dict(row) for row in resume_layout_signal_rule_rows],
        "resume_profiles": [dict(row) for row in resume_profile_rows],
        "archetype_resume_configs": {row["archetype_key"]: dict(row) for row in archetype_resume_config_rows},
        "resume_generation_settings": {row["key"]: row["value"] for row in resume_generation_setting_rows},
        "resume_profile_signal_rules": [dict(row) for row in resume_profile_signal_rule_rows],
        "role_records": [dict(row) for row in role_rows],
        "role_requirements": [dict(row) for row in role_requirement_rows],
        "role_archetype_scores": [dict(row) for row in role_archetype_score_rows],
    }


def normalize_requested_archetype(value):
    normalized = normalize_resume_profile_key(value)
    return ARCHETYPE_ALIASES.get(normalized, normalized)


def build_archetype_config(archetype_name, db_data=None):
    import json as _json
    archetype_name = normalize_requested_archetype(archetype_name)
    settings = (db_data or {}).get("resume_generation_settings") or {}
    defaults_raw = settings.get("defaults", "{}")
    try:
        defaults = _json.loads(defaults_raw) if isinstance(defaults_raw, str) else dict(defaults_raw or {})
    except Exception:
        defaults = {}

    db_configs = (db_data or {}).get("archetype_resume_configs") or {}
    db_row = db_configs.get(archetype_name)
    if not db_row:
        raise SystemExit(
            f"No DB-backed resume archetype config for '{archetype_name}'. "
            "Run: node scripts/db/migrate-archetype-configs.mjs"
        )

    def _parse(field):
        val = db_row.get(field)
        if isinstance(val, str):
            try:
                return _json.loads(val)
            except Exception:
                return {}
        return val or {}

    caution_rules = _parse("caution_rules")
    entry = {
        "id": db_row.get("id", archetype_name),
        "label": db_row.get("label", ""),
        "caution_rules": caution_rules if isinstance(caution_rules, list) else [],
        "experience": _parse("experience"),
        "skills": _parse("skills"),
        "education": _parse("education"),
        "certifications": _parse("certifications"),
    }
    entry["_meta"] = {
        "template_key": db_row.get("template_key") or archetype_name,
        "subtitle": db_row.get("subtitle") or "Software Engineer",
        "summary": db_row.get("summary") or "",
        "defaults": defaults,
    }
    return entry


def infer_relevance_context(archetype, keywords):
    keyword_text = " ".join(normalize_text(item) for item in keywords if item)
    archetype_id = archetype.get("id", "")
    return {
        "security": any(token in keyword_text for token in ["security", "cyber", "infosec", "network", "tcp/ip", "routing", "switching", "windows server", "linux"]),
        "xr": any(token in keyword_text for token in ["vr", "xr", "ar", "quest", "vive", "psvr", "hololens", "spatial", "immersive"]),
        "game": any(token in keyword_text for token in ["unity", "game", "gameplay", "liveops", "live service", "multiplayer"]),
        "backend": any(token in keyword_text for token in ["backend", "api", "apis", "redis", "azure", "sql", "cloud", "playfab"]),
    }


def split_secondary_categories(value):
    return {part.strip() for part in str(value or "").split("|") if part.strip()}


def skill_display_text(skill):
    return str(skill.get("resume_display") or skill["skill_name"]).strip()


def default_render_group(skill):
    category = skill["category"]
    if category == "language":
        return "programming"
    if category == "engine":
        return "engines"
    if category in {"backend", "cloud"}:
        return "backend-cloud"
    if category == "networking":
        return "networking"
    if category == "tool":
        return "tools"
    if category == "platform":
        return "xr-platforms"
    if category == "security":
        return "security"
    return category


def build_skill_rule_index(db_data, archetype_name):
    rules = {}
    for row in db_data.get("skill_resume_rules", []):
        if normalize_text(row.get("profile_key")) != normalize_text(archetype_name):
            continue
        if not parse_bool(row.get("approved"), True):
            continue
        rules[row.get("skill_normalized")] = dict(row)
    return rules


def build_group_rule_index(db_data, archetype_name):
    rules = {}
    for row in db_data.get("resume_group_rules", []):
        if normalize_text(row.get("profile_key")) != normalize_text(archetype_name):
            continue
        if not parse_bool(row.get("approved"), True):
            continue
        group_id = row.get("group_id")
        if not group_id:
            continue
        entry = rules.setdefault(group_id, {"group_id": group_id})
        entry.update(dict(row))
    return rules


def safety_group_rule(group_id):
    label_overrides = {
        "programming": "Programming",
        "engines": "Engines",
        "backend-cloud": "Backend & Cloud",
        "networking": "Networking",
        "xr-platforms": "Platforms",
        "security": "Security",
        "tools": "Tools",
    }
    return {
        "group_id": group_id,
        "label": label_overrides.get(group_id, group_id.replace("-", " ").title()),
        "group_rank": 999,
        "min_items_standalone": 1,
        "singleton_merge_target": None,
        "max_items": 8,
        "approved": 1,
        "_safety_fallback": True,
    }


def skill_rule_value(skill, rule, key, default=None):
    if rule and rule.get(key) is not None:
        return rule.get(key)
    return default


def build_skill_alias_index(db_data):
    aliases = {}
    for row in db_data.get("skill_aliases", []):
        skill_key = normalize_text(row.get("skill_normalized"))
        if not skill_key:
            continue
        entry = aliases.setdefault(skill_key, set())
        for value in (row.get("alias_name"), row.get("alias_normalized")):
            normalized = normalize_text(value)
            if normalized:
                entry.add(normalized)
    return aliases


def direct_match_strength(skill, keyword_set, alias_index):
    names = {
        normalize_text(skill.get("skill_name")),
        normalize_text(skill.get("skill_normalized")),
    }
    names.update(alias_index.get(normalize_text(skill.get("skill_normalized") or skill.get("skill_name")), set()))
    names.discard("")
    if not names or not keyword_set:
        return 0

    strength = 0
    for keyword in keyword_set:
        if keyword in names:
            strength = max(strength, 3)
            continue
        if any(keyword in name or name in keyword for name in names):
            strength = max(strength, 2)
    return strength


def normalize_role_context_value(value):
    text = str(value or "").strip()
    return text or None


def build_role_record_index(db_data):
    index = {}
    for row in db_data.get("role_records", []):
        company = normalize_text(row.get("company"))
        role = normalize_text(row.get("role"))
        if not company or not role:
            continue
        index[(company, role)] = dict(row)
    return index


def build_role_requirements_index(db_data):
    index = {}
    for row in db_data.get("role_requirements", []):
        role_id = row.get("role_id")
        if role_id is None:
            continue
        index.setdefault(role_id, []).append(dict(row))
    return index


def normalize_archetype_score_key(value):
    normalized = normalize_resume_profile_key(value)
    aliases = {
        "xr": "vr",
        "security": "cyber",
        "forensics": "forensic",
    }
    return aliases.get(normalized, normalized)


def build_role_archetype_score_index(db_data):
    index = {}
    for row in db_data.get("role_archetype_scores", []):
        if not parse_bool(row.get("approved"), True):
            continue
        role_id = row.get("role_id")
        normalized_key = normalize_archetype_score_key(row.get("archetype_key"))
        if role_id is None or not normalized_key:
            continue
        index.setdefault(role_id, []).append(
            {
                **dict(row),
                "archetype_key": normalized_key,
                "score": float(row.get("score") or 0.0),
                "rank": parse_int(row.get("rank"), 999),
            }
        )

    for role_id, rows in index.items():
        rows.sort(key=lambda item: (item["rank"], -item["score"], item["archetype_key"]))
    return index


def build_role_archetype_score_map(rows):
    score_map = {}
    for row in rows or []:
        key = normalize_archetype_score_key(row.get("archetype_key"))
        if not key:
            continue
        score = float(row.get("score") or 0.0)
        score_map[key] = max(score_map.get(key, 0.0), score)
    return score_map


def resolve_active_role_record(db_data, role_context):
    if not role_context:
        return None
    company = normalize_text(role_context.get("company"))
    role = normalize_text(role_context.get("role"))
    if not company or not role:
        return None
    return build_role_record_index(db_data).get((company, role))


def resolve_role_archetype_scores(db_data, role_context):
    matched_role = resolve_active_role_record(db_data, role_context)
    if not matched_role:
        return []
    role_id = row_value(matched_role, "num", "id", "role_id")
    if role_id is None:
        return []
    return list(build_role_archetype_score_index(db_data).get(role_id, []))


def enrich_role_archetype_score_map(role_archetype_map, archetype_name, archetype):
    enriched = dict(role_archetype_map or {})
    explicit_candidates = {
        normalize_archetype_score_key(archetype_name),
        normalize_archetype_score_key(archetype.get("id", "")),
        normalize_archetype_score_key(archetype.get("label", "")),
        normalize_archetype_score_key(ARCHETYPE_ID_MAP.get(archetype_name, "")),
    }
    explicit_candidates.discard("")
    for candidate in explicit_candidates:
        enriched[candidate] = max(enriched.get(candidate, 0.0), 0.35)
    return enriched


def resolve_role_context(db_data, args):
    role_context = {
        "role": normalize_role_context_value(getattr(args, "job_role", None)),
        "company": normalize_role_context_value(getattr(args, "job_company", None)),
        "notes": normalize_role_context_value(getattr(args, "job_notes", None)),
        "location_text": normalize_role_context_value(getattr(args, "job_location", None)),
        "work_model": normalize_role_context_value(getattr(args, "job_work_model", None)),
        "compensation_text": normalize_role_context_value(getattr(args, "job_compensation", None)),
    }
    if role_context["company"] and role_context["role"]:
        matched = build_role_record_index(db_data).get(
            (normalize_text(role_context["company"]), normalize_text(role_context["role"]))
        )
        if matched:
            role_context["role_id"] = row_value(matched, "num", "id", "role_id")
            role_context["notes"] = role_context["notes"] or normalize_role_context_value(matched.get("notes"))
            role_context["location_text"] = role_context["location_text"] or normalize_role_context_value(matched.get("location_text"))
            role_context["work_model"] = role_context["work_model"] or normalize_role_context_value(matched.get("work_model"))
            role_context["compensation_text"] = role_context["compensation_text"] or normalize_role_context_value(matched.get("compensation_text"))
    return role_context


NON_US_LOCATION_HINTS = {
    "canada", "vancouver", "victoria", "montreal", "toronto", "british columbia",
    "romania", "bucharest",
    "united kingdom", "uk", "england", "manchester", "birmingham",
    "sweden", "stockholm",
    "colombia", "bogota",
    "china", "shanghai",
    "korea", "south korea", "seoul",
    "portugal", "brazil",
}

US_LOCATION_HINTS = {
    "united states", "usa", "u s a", "us remote", "remote us",
    "california", "texas", "florida", "colorado", "alabama", "wisconsin",
    "orlando", "austin", "los angeles", "la", "madison", "redwood city", "huntsville", "colorado springs",
}


def role_target_is_non_us(role_context):
    location = normalize_text(role_context.get("location_text"))
    work_model = normalize_text(role_context.get("work_model"))
    compensation = normalize_text(role_context.get("compensation_text"))
    combined = " ".join(part for part in [location, work_model, compensation] if part).strip()
    if not combined:
        return False
    if any(hint in combined for hint in NON_US_LOCATION_HINTS):
        return True
    if " cad" in f" {combined} " or " eur" in f" {combined} " or " gbp" in f" {combined} ":
        return True
    if any(hint in combined for hint in US_LOCATION_HINTS):
        return False
    if re.search(r"\b[a-z]{2}\b", location) and not re.search(r"\b(?:tx|fl|al|co|wi)\b", location):
        return True
    return False


def build_runtime_feature_map(db_data, role_context, keywords, roles, education_entries, certification_entries):
    keyword_tokens = {normalize_text(item) for item in keywords if normalize_text(item)}
    role_notes = normalize_text(role_context.get("notes"))
    combined_context = " ".join(
        normalize_text(role_context.get(key))
        for key in ["role", "company", "notes", "location_text", "work_model", "compensation_text"]
        if role_context.get(key)
    ).strip()

    matched_role = resolve_active_role_record(db_data, role_context)

    requirements = []
    if matched_role:
        requirements = build_role_requirements_index(db_data).get(matched_role.get("num"), [])
    role_archetype_scores = resolve_role_archetype_scores(db_data, role_context)
    role_archetype_map = build_role_archetype_score_map(role_archetype_scores)

    priority_totals = {"required": 0, "preferred": 0, "bonus": 0, "unknown": 0}
    priority_matched = {"required": 0, "preferred": 0, "bonus": 0, "unknown": 0}
    cluster_scores = {
        "backend": 0,
        "gameplay": 0,
        "liveops": 0,
        "xr": 0,
        "security": 0,
        "forensics": 0,
        "networking": 0,
        "frontend": 0,
        "systems": 0,
        "certifications": 0,
        "education": 0,
    }
    direct_match_count = 0
    unmatched_count = 0
    weak_context_count = 0

    cluster_tokens = {
        "backend": ["backend", "api", "redis", "sql", "cloud", "azure", "playfab", "distributed"],
        "gameplay": ["gameplay", "player", "combat", "progression", "unity", "game systems"],
        "liveops": ["liveops", "live ops", "live-service", "live service", "events", "monetization", "economy"],
        "xr": ["vr", "xr", "augmented reality", "mixed reality", "quest", "vive", "hololens", "psvr", "spatial", "immersive"],
        "security": ["security", "cyber", "infosec", "network", "ccna", "clearance", "secret", "dod"],
        "forensics": ["forensic", "forensics", "incident response", "investigation", "malware", "soc"],
        "networking": ["network", "networking", "tcp/ip", "routing", "switching", "cisco", "ccna"],
        "frontend": ["frontend", "front end", "react", "javascript", "typescript", "html", "css"],
        "systems": ["systems", "sysml", "cameo", "mbse", "dodaf", "uaf", "model based"],
        "certifications": ["cert", "certificate", "certification", "ccna", "ocsmp", "clearance"],
        "education": ["degree", "b.s.", "bachelor", "college", "education", "coursework", "associate"],
    }

    for requirement in requirements:
        priority = normalize_text(requirement.get("priority")) or "unknown"
        if priority not in priority_totals:
            priority = "unknown"
        priority_totals[priority] += 1

        requirement_text = " ".join(
            [
                normalize_text(requirement.get("requirement_name")),
                normalize_text(requirement.get("requirement_normalized")),
                normalize_text(requirement.get("raw_text")),
                normalize_text(requirement.get("notes")),
            ]
        ).strip()
        matched = normalize_text(requirement.get("matched_entity_type")) in {"skill", "capability"} or normalize_text(requirement.get("match_method")) in {"exact", "alias", "similar"}
        if matched:
            priority_matched[priority] += 1
            direct_match_count += 1
        else:
            unmatched_count += 1

        if normalize_text(requirement.get("match_method")) in {"similar"}:
            weak_context_count += 1
        if normalize_text(requirement.get("match_method")) in {"new_candidate", "unmatched"}:
            weak_context_count += 2

        for cluster, tokens in cluster_tokens.items():
            if any(token in requirement_text for token in tokens):
                cluster_scores[cluster] += 3 if matched else 1

    for token in keyword_tokens:
        if len(token) < 3:
            continue
        for cluster, tokens in cluster_tokens.items():
            if any(fragment in token or token in fragment for fragment in tokens):
                cluster_scores[cluster] += 2

    if "cerritos" in combined_context or "certificate" in combined_context or "ccna" in combined_context:
        cluster_scores["education"] += 2
        cluster_scores["certifications"] += 2

    cert_sensitivity = cluster_scores["certifications"] + cluster_scores["security"] + cluster_scores["networking"] + cluster_scores["forensics"]
    education_sensitivity = cluster_scores["education"] + cluster_scores["systems"]
    cerritos_fit = 0
    if education_entries and _SUPPLEMENTAL_SCHOOL and any(item.get("school") == _SUPPLEMENTAL_SCHOOL for item in education_entries):
        cerritos_fit += 2
    if any(token in combined_context for token in ["ccna", "network", "networking", "cyber", "forensic", "forensics"]):
        cerritos_fit += 4
    if cluster_scores["frontend"] > 0 and cert_sensitivity == 0:
        cerritos_fit -= 2

    cert_request_tokens = ["certification", "certifications", "certificate", "cert", "ccna", "ocsmp", "clearance", "security+", "cissp"]
    cert_requested = any(token in combined_context for token in cert_request_tokens)
    matched_cert_requirement_count = 0
    for requirement in requirements:
        requirement_text = " ".join(
            [
                normalize_text(requirement.get("requirement_name")),
                normalize_text(requirement.get("requirement_normalized")),
                normalize_text(requirement.get("raw_text")),
                normalize_text(requirement.get("notes")),
            ]
        )
        if any(token in requirement_text for token in cert_request_tokens) and (
            normalize_text(requirement.get("matched_entity_type")) in {"skill", "capability"}
            or normalize_text(requirement.get("match_method")) in {"exact", "alias", "similar"}
        ):
            matched_cert_requirement_count += 1

    total_requirements = sum(priority_totals.values())
    total_matched = sum(priority_matched.values())
    matched_ratio = round(total_matched / total_requirements, 3) if total_requirements else 0.0
    unmatched_ratio = round(unmatched_count / total_requirements, 3) if total_requirements else 0.0
    direct_match_density = round(direct_match_count / max(1, total_requirements + len(keyword_tokens)), 3)

    experience_bullet_count = 0
    context_only_bullet_count = 0
    for role in roles:
        for bullet in role.get("direct_bullets", []):
            experience_bullet_count += 1
            if normalize_text(bullet.get("inclusion_mode", "context")) == "context":
                context_only_bullet_count += 1
        for project in role.get("projects", []):
            for bullet in project.get("bullets", []):
                experience_bullet_count += 1
                if normalize_text(bullet.get("inclusion_mode", "context")) == "context":
                    context_only_bullet_count += 1

    weak_context_pressure = round(
        (weak_context_count + context_only_bullet_count) / max(1, total_requirements + experience_bullet_count),
        3,
    )

    return {
        "role_id": matched_role.get("num") if matched_role else None,
        "role_archetype_scores": role_archetype_scores,
        "role_archetype_score_map": role_archetype_map,
        "role_primary_archetype": role_archetype_scores[0]["archetype_key"] if role_archetype_scores else None,
        "keywords": sorted(keyword_tokens),
        "role_notes": role_notes,
        "requirement_counts": {
            "total": total_requirements,
            "matched": total_matched,
            "unmatched": unmatched_count,
            "by_priority_total": priority_totals,
            "by_priority_matched": priority_matched,
        },
        "matched_requirement_ratio": matched_ratio,
        "unmatched_requirement_ratio": unmatched_ratio,
        "cluster_scores": cluster_scores,
        "cert_sensitivity": cert_sensitivity,
        "education_sensitivity": education_sensitivity,
        "cerritos_fit": cerritos_fit,
        "cert_requested": cert_requested,
        "matched_cert_requirement_count": matched_cert_requirement_count,
        "direct_match_density": direct_match_density,
        "weak_context_pressure": weak_context_pressure,
        "is_games": (
            cluster_scores.get("gameplay", 0) > 0
            or cluster_scores.get("liveops", 0) > 0
            or cluster_scores.get("xr", 0) > 0
        ),
    }


def evaluate_trigger_condition(condition, role_context):
    normalized = normalize_text(condition)
    if not normalized:
        return True
    if normalized == "role_not_us":
        return role_target_is_non_us(role_context)
    if normalized == "is_games":
        return bool(role_context.get("is_games"))
    return False


def should_include_skill(skill, rule, group_id, pinned_names, match_strength, context, role_context):
    level = normalize_text(skill["level"])
    secondary = split_secondary_categories(skill.get("secondary_categories"))
    visibility = normalize_text(skill_rule_value(skill, rule, "visibility", "show")) or "show"
    trigger_condition = skill_rule_value(skill, rule, "trigger_condition")
    include_default = parse_bool(skill.get("include_default"), True)
    require_direct_match = parse_bool(skill.get("require_direct_match"), False)
    has_direct_match = match_strength > 0

    if visibility == "hidden":
        return False
    if normalize_text(skill.get("resume_visibility")) == "hidden":
        return False

    # Never include level=none or level=exposure skills unless explicitly pinned in archetype config.
    if level in {"none", "exposure"} and skill["skill_name"] not in pinned_names:
        return False

    # GDScript only if gdscript is in keywords; Godot is paired via alias so both appear together
    if normalize_text(skill.get("skill_name", "")) in {"gdscript", "gd script"}:
        if not has_direct_match:
            return False

    if skill["skill_name"] in pinned_names:
        return True
    if has_direct_match:
        return True
    if not include_default and rule is None:
        return False
    if trigger_condition and not evaluate_trigger_condition(trigger_condition, role_context):
        return False
    if require_direct_match:
        return False
    if visibility == "context" and not include_default:
        return False

    if group_id == "programming":
        if level in {"exposure", "none"}:
            return False
        if level in {"expert", "advanced"}:
            return True
        if visibility == "show":
            return True
        return level == "intermediate" and not secondary.intersection({"Roblox", "Platforms & Engines"})
    if group_id == "engines":
        if visibility == "show":
            return True
        return context["game"] or context["xr"] or level in {"expert", "advanced"}
    if group_id == "backend-cloud":
        return visibility == "show" or level in {"expert", "advanced", "intermediate"}
    if group_id == "networking":
        if not (context["backend"] or context["security"]):
            return False
        return visibility == "show" or level in {"expert", "advanced", "intermediate"}
    if group_id == "xr-platforms":
        return visibility == "show" or context["xr"]
    if group_id == "security":
        return visibility == "show" or context["security"]
    if group_id == "tools":
        if secondary.intersection({"Creative Tools", "Roblox"}):
            return False
        if level in {"expert", "advanced"}:
            return True
        if level == "intermediate":
            return secondary.intersection({"Developer Tools", "Cloud / Backend"}) or context["backend"] or visibility == "show"
        return visibility == "show"
    return visibility == "show" or level in {"expert", "advanced", "intermediate"}


def build_skill_groups(skills, archetype, keywords):
    fail("build_skill_groups now requires db_data and archetype_name")


def build_skill_selection_policy(archetype, runtime_feature_map, max_pages):
    defaults = archetype.get("_meta", {}).get("defaults", {})
    configured_groups = int(defaults.get("max_skill_groups", 6))
    direct_match_density = float((runtime_feature_map or {}).get("direct_match_density") or 0.0)
    weak_context_pressure = float((runtime_feature_map or {}).get("weak_context_pressure") or 0.0)
    is_games_role = bool((runtime_feature_map or {}).get("is_games"))
    low_fit = direct_match_density < 0.2 or weak_context_pressure >= 0.45
    high_pressure = max_pages <= 1 or weak_context_pressure >= 0.55
    max_groups = configured_groups
    if low_fit:
        max_groups = max(3, configured_groups - 1)
    if high_pressure:
        max_groups = max(2, max_groups - 1)
    if is_games_role and direct_match_density >= 0.35:
        max_groups = max(4, max_groups)
    max_groups = min(5, max_groups)
    return {
        "configured_max_groups": configured_groups,
        "max_groups": max_groups,
        "direct_match_density": direct_match_density,
        "weak_context_pressure": weak_context_pressure,
        "low_fit_role": low_fit,
        "layout_compaction": high_pressure,
        "shared_signal_weight": 8 if low_fit else 5,
        "is_games_role": is_games_role,
    }


def group_signal_score(group_id, runtime_feature_map):
    cluster_scores = (runtime_feature_map or {}).get("cluster_scores") or {}
    mapping = {
        "programming": ("backend", "systems", "frontend"),
        "engines": ("gameplay", "graphics", "systems"),
        "backend-cloud": ("backend", "cloud", "systems"),
        "networking": ("networking", "security", "systems"),
        "xr-platforms": ("xr", "gameplay"),
        "security": ("security", "forensics", "networking"),
        # Tool-heavy game/graphics roles often surface workflow, renderer, and engine-adjacent
        # requirements that should compete with backend/cloud under tight one-page budgets.
        "tools": ("backend", "systems", "graphics", "frontend", "gameplay", "xr"),
    }
    clusters = mapping.get(group_id, ())
    return sum(parse_int(cluster_scores.get(cluster), 0) for cluster in clusters)


def build_skill_groups(db_data, skills, archetype_name, archetype, keywords, role_context=None, runtime_feature_map=None, max_pages=1):
    level_weight = {"expert": 5, "advanced": 4, "intermediate": 3, "basic": 2, "exposure": 1}
    pinned_names = {SKILL_ID_MAP.get(item, item) for item in archetype.get("skills", {}).get("pinned_items", [])}
    keyword_set = {normalize_text(item) for item in keywords if normalize_text(item)}
    context = infer_relevance_context(archetype, keywords)
    role_context = role_context or {}
    runtime_feature_map = runtime_feature_map or {}
    selection_policy = build_skill_selection_policy(archetype, runtime_feature_map, max_pages)
    max_groups = selection_policy["max_groups"]
    skill_rules = build_skill_rule_index(db_data, archetype_name)
    group_rules = build_group_rule_index(db_data, archetype_name)
    alias_index = build_skill_alias_index(db_data)
    use_safety_group_rules = not group_rules

    _game_archetypes = {"gameplay", "liveops", "vr"}
    _is_games_jd = bool(runtime_feature_map.get("is_games"))
    _is_games_archetype = archetype_name in _game_archetypes
    role_context = {**role_context, "is_games": _is_games_jd or _is_games_archetype, "archetype_name": archetype_name}

    grouped = {}
    for skill in skills:
        rule = skill_rules.get(skill.get("skill_normalized") or normalize_text(skill["skill_name"]))
        group_id = normalize_text(skill_rule_value(skill, rule, "render_group", default_render_group(skill))) or default_render_group(skill)
        match_strength = direct_match_strength(skill, keyword_set, alias_index)
        if group_id not in group_rules:
            if not use_safety_group_rules:
                continue
            group_rules[group_id] = safety_group_rule(group_id)
        if not should_include_skill(skill, rule, group_id, pinned_names, match_strength, context, role_context):
            continue

        score = level_weight.get(normalize_text(skill["level"]), 0)
        score += parse_int(skill.get("resume_priority"), 0)
        score += parse_int(skill_rule_value(skill, rule, "item_rank", 0), 0)
        score += min(30, group_signal_score(group_id, runtime_feature_map) * int(selection_policy["shared_signal_weight"]))
        if skill["skill_name"] in pinned_names:
            score += 20
        if match_strength:
            score += (match_strength * 10) + parse_int(skill_rule_value(skill, rule, "direct_match_boost", 0), 0)
        elif selection_policy["low_fit_role"] and not parse_bool(skill.get("require_direct_match"), False):
            score -= 8

        item = {
            "score": score,
            "name": skill["skill_name"],
            "display": skill_display_text(skill),
            "emphasis": normalize_text(skill_rule_value(skill, rule, "emphasis", "plain")) or "plain",
            "item_rank": parse_int(skill_rule_value(skill, rule, "item_rank", 0), 0),
            "direct_match_boost": parse_int(skill_rule_value(skill, rule, "direct_match_boost", 0), 0),
            "singleton_penalty": parse_int(skill_rule_value(skill, rule, "singleton_penalty", 0), 0),
            "matched_keyword": match_strength > 0,
            "direct_match_strength": match_strength,
            "visibility": visibility if (visibility := normalize_text(skill_rule_value(skill, rule, "visibility", "show")) or "show") else "show",
            "applied_skill_rule": dict(rule) if rule else None,
            "source_group_id": group_id,
        }
        entry = grouped.setdefault(
            group_id,
            {
                "group_id": group_id,
                "label": row_value(group_rules[group_id], "label", default=group_id.title()),
                "group_rank": parse_int(row_value(group_rules[group_id], "group_rank"), 0),
                "min_items_standalone": parse_int(row_value(group_rules[group_id], "min_items_standalone"), 1),
                "singleton_merge_target": row_value(group_rules[group_id], "singleton_merge_target"),
                "max_items": parse_int(row_value(group_rules[group_id], "max_items"), 8),
                "group_rule": dict(group_rules[group_id]),
                "merged_from": [],
                "items": [],
            },
        )
        entry["items"].append(item)

    for entry in grouped.values():
        entry["items"].sort(
            key=lambda item: (
                -item["direct_match_strength"],
                -item["direct_match_boost"],
                -item["item_rank"],
                -item["score"],
                item["display"].lower(),
            )
        )

    for group_id, entry in list(grouped.items()):
        if len(entry["items"]) >= entry["min_items_standalone"]:
            continue
        target_id = entry.get("singleton_merge_target")
        if not target_id or target_id not in grouped:
            continue
        grouped[target_id]["merged_from"].append(
            {
                "group_id": entry["group_id"],
                "label": entry["label"],
                "reason": "singleton_merge",
            }
        )
        for item in entry["items"]:
            item["score"] -= item.get("singleton_penalty", 0)
            item["merged_from_group"] = entry["group_id"]
            grouped[target_id]["items"].append(item)
        del grouped[group_id]

    rendered = []
    for entry in grouped.values():
        entry["items"].sort(
            key=lambda item: (
                -item["direct_match_strength"],
                -item["direct_match_boost"],
                -item["item_rank"],
                -item["score"],
                item["display"].lower(),
            )
        )
        entry["items"] = entry["items"][: entry["max_items"]]
        if not entry["items"]:
            continue
        if selection_policy["low_fit_role"]:
            strong_items = [item for item in entry["items"] if item.get("matched_keyword") or item.get("direct_match_strength", 0) > 0]
            if strong_items:
                entry["items"] = strong_items + [item for item in entry["items"] if item not in strong_items][: max(0, entry["max_items"] - len(strong_items))]
        if len(entry["items"]) == 1:
            item = entry["items"][0]
            if item.get("visibility") == "context" and not item.get("matched_keyword") and item.get("score", 0) < 15:
                continue
        entry["score"] = max(item["score"] for item in entry["items"])
        if selection_policy["low_fit_role"] and len(entry["items"]) > 1:
            matched_count = sum(1 for item in entry["items"] if item.get("matched_keyword"))
            if matched_count == 0:
                defaultish_count = sum(1 for item in entry["items"] if item.get("visibility") != "show")
                if defaultish_count >= len(entry["items"]) or entry["score"] < 22:
                    continue
        entry["selection_policy"] = dict(selection_policy)
        entry["shared_signal_score"] = group_signal_score(entry["group_id"], runtime_feature_map)
        rendered.append(entry)

    def sort_key(entry):
        matched_count = sum(1 for item in entry["items"] if item.get("matched_keyword"))
        direct_strength = sum(parse_int(item.get("direct_match_strength"), 0) for item in entry["items"])
        effective_rank = entry["group_rank"]
        if selection_policy["low_fit_role"] and matched_count:
            shared_signal = parse_int(entry.get("shared_signal_score"), 0)
            effective_rank -= min(40, (matched_count * 12) + (direct_strength * 3) + (shared_signal * 2))
        return (effective_rank, entry["group_rank"], -entry["score"], entry["label"].lower())

    rendered.sort(key=sort_key)

    # Combine backend-cloud + networking into one line if all items fit within the char limit
    bc_idx = next((i for i, e in enumerate(rendered) if e["group_id"] == "backend-cloud"), None)
    net_idx = next((i for i, e in enumerate(rendered) if e["group_id"] == "networking"), None)
    if bc_idx is not None and net_idx is not None:
        bc_entry = rendered[bc_idx]
        net_entry = rendered[net_idx]
        combined_items = bc_entry["items"] + net_entry["items"]
        combined_label = "Backend & Networking"
        combined_formatted = format_skill_group_items(combined_items)
        if len(combined_label) + len(": ") + len(combined_formatted) <= SKILL_GROUP_CHAR_LIMIT:
            bc_entry = dict(bc_entry)
            bc_entry["items"] = combined_items
            bc_entry["label"] = combined_label
            bc_entry["group_id"] = "backend-networking"
            rendered = [e for i, e in enumerate(rendered) if i != bc_idx and i != net_idx]
            rendered.insert(bc_idx, bc_entry)

    if selection_policy["layout_compaction"]:
        compacted = []
        for entry in rendered:
            compacted_entry = dict(entry)
            compacted_entry["items"] = list(entry["items"])
            while len(compacted_entry["items"]) > 2 and not any(item.get("matched_keyword") for item in compacted_entry["items"][-1:]):
                last = compacted_entry["items"][-1]
                if last.get("visibility") == "show" or last.get("emphasis") == "primary":
                    break
                compacted_entry["items"].pop()
            compacted.append(compacted_entry)
        rendered = compacted
    return rendered[:max_groups]


def project_config(project_id):
    if project_id not in LEGACY_PROJECTS:
        fail(f"Missing project mapping for {project_id}")
    return LEGACY_PROJECTS[project_id]


def identify_point_project(role_id, bullet_text):
    normalized = normalize_text(bullet_text)
    for project_id, project in LEGACY_PROJECTS.items():
        if project["role"] != role_id:
            continue
        if any(normalize_text(keyword) in normalized for keyword in project["keywords"]):
            return project_id
    return None


def build_resume_points(db_data):
    points = []
    for bullet in db_data["bullets"]:
        company = bullet["company"]
        title = bullet["title"]
        role_id = None
        for candidate_role_id, role_key in ROLE_ID_MAP.items():
            if role_key == (company, title):
                role_id = candidate_role_id
                break
        if role_id is None:
            continue
        project_id = identify_point_project(role_id, bullet["text"])
        points.append(
            {
                "role_id": role_id,
                "experience_id": bullet["exp_id"],
                "project_id": project_id,
                "text": bullet["text"].strip(),
                "dedupe_key": normalize_text(bullet["text"]),
                "tags": [tag.strip() for tag in str(bullet.get("tags") or "").split("|") if tag.strip()],
                "strength": int(bullet.get("strength") or 0),
                "company": company,
                "title": title,
                "start_date": bullet["start_date"],
                "end_date": bullet["end_date"],
            }
        )
    return points


def score_resume_point(point, preferred_tags, pinned_fragments, keywords):
    text = normalize_text(point["text"])
    score = point["strength"] * 10
    score += sum(6 for tag in point["tags"] if tag in preferred_tags)
    score += sum(5 for keyword in keywords if keyword and keyword in text)
    pinned = any(fragment in text for fragment in pinned_fragments)
    if pinned:
        score += 1000
    return {
        **point,
        "score": score,
        "pinned": pinned,
    }


def has_resume_schema(db_data):
    required = {"projects", "resume_points"}
    return required.issubset(set(db_data.get("available_tables") or set()))


def role_lookup_by_id(db_data):
    lookup = {}
    for role_id, role_key in ROLE_ID_MAP.items():
        company, title = role_key
        for row in db_data["experience"]:
            if row["company"] == company and row["title"] == title:
                lookup[role_id] = dict(row)
                break
    return lookup


def build_resume_project_index(db_data):
    project_skills = {}
    for row in db_data.get("project_skills", []):
        project_id = row_value(row, "project_id")
        if project_id is None:
            continue
        if not parse_bool(row_value(row, "approved", "is_approved"), True):
            continue
        project_skills.setdefault(project_id, []).append(
            {
                "name": row_value(row, "skill_name", "display_name", "skill_normalized", "name", default=""),
                "sort_priority": parse_int(row_value(row, "sort_priority", "priority", "display_order"), 0),
            }
        )

    projects = {}
    for row in db_data.get("projects", []):
        project_id = row_value(row, "id", "project_id")
        if project_id is None:
            continue
        if not parse_bool(row_value(row, "approved", "is_approved"), True):
            continue
        experience_id = row_value(row, "experience_id", "exp_id")
        stack = row_value(row, "stack", "tech_stack", "stack_text", default="")
        if not str(stack).strip():
            skills = sorted(
                project_skills.get(project_id, []),
                key=lambda item: (item["sort_priority"], normalize_text(item["name"])),
            )
            stack = ", ".join(item["name"] for item in skills if str(item["name"]).strip())
        projects[project_id] = {
            "id": project_id,
            "experience_id": experience_id,
            "label": row_value(row, "name", "label", "project_name", "title", default=""),
            "stack": stack,
            "sort_priority": parse_int(row_value(row, "sort_priority", "priority", "display_order"), 0),
        }
    return projects


def normalize_resume_profile_key(value):
    text = normalize_text(value)
    return re.sub(r"[^a-z0-9]+", "-", text).strip("-")


def default_profile_key_for_archetype(archetype_name):
    normalized = normalize_requested_archetype(archetype_name)
    return normalized


def build_resume_profile_index(db_data):
    return {
        normalize_resume_profile_key(row_value(row, "profile_key", default="")): dict(row)
        for row in db_data.get("resume_profiles", [])
        if parse_bool(row_value(row, "approved", "is_approved"), True)
        and normalize_resume_profile_key(row_value(row, "profile_key", default=""))
    }


def build_resume_profile_signals(db_data, archetype_name, keywords, role_context):
    keywords = [normalize_text(keyword) for keyword in keywords if normalize_text(keyword)]
    combined_context = " ".join(
        normalize_text(role_context.get(key))
        for key in ["role", "company", "location_text", "work_model", "compensation_text", "notes"]
        if role_context.get(key)
    )

    active_role = resolve_active_role_record(db_data, role_context)

    role_requirements = []
    if active_role:
        role_id = row_value(active_role, "num", "id", "role_id")
        role_requirements = [row for row in db_data.get("role_requirements", []) if row_value(row, "role_id") == role_id]
    role_archetype_scores = resolve_role_archetype_scores(db_data, role_context)
    role_archetype_map = build_role_archetype_score_map(role_archetype_scores)

    matched_count = sum(1 for row in role_requirements if normalize_text(row_value(row, "match_method", default="unmatched")) not in {"unmatched", "new-candidate", "new_candidate"})
    unmatched_count = sum(1 for row in role_requirements if normalize_text(row_value(row, "match_method", default="unmatched")) in {"unmatched", "new-candidate", "new_candidate"})
    total_count = len(role_requirements)
    requirement_match_ratio = (matched_count / total_count) if total_count else 0.0

    security_tokens = ["security", "cyber", "infosec", "forensic", "forensics", "threat", "incident", "vulnerability", "network", "tcp/ip", "routing", "switching", "windows server", "linux", "siem", "soc", "detection"]
    backend_tokens = ["backend", "api", "apis", "redis", "azure", "sql", "database", "cloud", "service", "distributed", "scalability", "microservices", "playfab"]
    gameplay_tokens = ["gameplay", "unity", "player", "combat", "mission", "physics", "progression", "game systems", "multiplayer"]
    liveops_tokens = ["liveops", "live ops", "live-ops", "live service", "economy", "monetization", "events", "retention", "engagement", "rewards"]
    vr_tokens = ["vr", "xr", "virtual reality", "mixed reality", "quest", "vive", "psvr", "hololens", "spatial", "immersive", "simulation"]

    def count_token_hits(tokens):
        text = " ".join(
            [
                combined_context,
                " ".join(normalize_text(row_value(req, "raw_text", "requirement_name", default="")) for req in role_requirements),
                " ".join(keywords),
            ]
        )
        return sum(1 for token in tokens if token and token in text)

    fit_band_direct = 1 if (matched_count >= 4 or requirement_match_ratio >= 0.55 or (matched_count >= 2 and total_count <= 3)) else 0
    fit_band_adjacent = 1 if not fit_band_direct and (matched_count >= 2 or requirement_match_ratio >= 0.25 or count_token_hits(backend_tokens + security_tokens + gameplay_tokens + vr_tokens + liveops_tokens) >= 2) else 0
    low_match_general_fallback = 1 if total_count > 0 and not fit_band_direct and not fit_band_adjacent else 0

    normalized_archetype = normalize_requested_archetype(archetype_name)
    signals = {
        "requirement_count_total": total_count,
        "requirement_count_matched": matched_count,
        "requirement_count_unmatched": unmatched_count,
        "requirement_match_ratio": round(requirement_match_ratio, 4),
        "jd_signal_security": count_token_hits(security_tokens),
        "jd_signal_backend": count_token_hits(backend_tokens),
        "jd_signal_gameplay": count_token_hits(gameplay_tokens),
        "jd_signal_liveops": count_token_hits(liveops_tokens),
        "jd_signal_vr": count_token_hits(vr_tokens),
        "fit_band_direct": fit_band_direct,
        "fit_band_adjacent": fit_band_adjacent,
        "low_match_general_fallback": low_match_general_fallback,
        "archetype_online": 1 if normalized_archetype == "gameserver" else 0,
    }
    signals[f"archetype_{normalized_archetype}"] = 1
    if role_archetype_scores:
        primary_key = normalize_resume_profile_key(role_archetype_scores[0].get("archetype_key", ""))
        if primary_key:
            signals[f"role_primary_archetype_is_{primary_key}"] = 1
    for archetype_key, score in role_archetype_map.items():
        signals[f"role_archetype_score_{archetype_key}"] = round(float(score), 4)
    if role_archetype_scores:
        signals["role_archetype_top_score"] = round(float(role_archetype_scores[0].get("score") or 0.0), 4)
    signals["role_archetype_scores"] = role_archetype_scores
    return signals


def profile_signal_rule_matches(rule, signals):
    signal_key = row_value(rule, "signal_key", default="")
    operator = normalize_text(row_value(rule, "operator", default="eq"))
    actual = signals.get(signal_key)
    threshold_numeric = row_value(rule, "threshold_numeric")
    threshold_text = row_value(rule, "threshold_text", default="")

    if operator == "present":
        return actual not in (None, "", 0)
    if operator == "absent":
        return actual in (None, "", 0)
    if actual is None:
        return False
    if operator == "eq":
        if threshold_numeric is not None:
            return float(actual) == float(threshold_numeric)
        return normalize_text(actual) == normalize_text(threshold_text)
    if threshold_numeric is None:
        return False
    if operator == "gte":
        return float(actual) >= float(threshold_numeric)
    if operator == "lte":
        return float(actual) <= float(threshold_numeric)
    return False


def score_resume_profiles(db_data, archetype_name, keywords, role_context):
    signals = build_resume_profile_signals(db_data, archetype_name, keywords, role_context)
    profiles_by_key = build_resume_profile_index(db_data)
    rules_by_profile = {}
    for row in db_data.get("resume_profile_signal_rules", []):
        if not parse_bool(row_value(row, "approved", "is_approved"), True):
            continue
        profile_key = normalize_resume_profile_key(row_value(row, "profile_key", default=""))
        if not profile_key or profile_key not in profiles_by_key:
            continue
        rules_by_profile.setdefault(profile_key, []).append(dict(row))

    scored = []
    default_profile_key = default_profile_key_for_archetype(archetype_name)
    role_archetype_scores = signals.get("role_archetype_scores") or []
    role_archetype_map = build_role_archetype_score_map(role_archetype_scores)
    for profile_key, profile in profiles_by_key.items():
        score = 0
        matched_rules = []
        blocked = False
        for rule in rules_by_profile.get(profile_key, []):
            if not profile_signal_rule_matches(rule, signals):
                continue
            weight = parse_int(row_value(rule, "weight"), 0)
            action = normalize_text(row_value(rule, "action", default="score")) or "score"
            matched_rules.append(
                {
                    "signal_key": row_value(rule, "signal_key", default=""),
                    "operator": row_value(rule, "operator", default=""),
                    "weight": weight,
                    "action": action,
                }
            )
            if action == "gate":
                blocked = True
            elif action == "penalty":
                score -= weight
            else:
                score += weight

        if profile_key == default_profile_key:
            score += 10
        if profile_key == "general":
            score += 5
        score += int(round(role_archetype_map.get(normalize_archetype_score_key(profile_key), 0.0) * 12))
        if profile_key == "cyber":
            score += int(round(role_archetype_map.get("forensic", 0.0) * 8))
        if profile_key == "forensic":
            score += int(round(role_archetype_map.get("cyber", 0.0) * 6))
        scored.append(
            {
                **dict(profile),
                "profile_key": profile_key,
                "score": score,
                "blocked": blocked,
                "matched_rules": matched_rules,
                "signals": signals,
            }
        )

    scored.sort(key=lambda item: (item["blocked"], -item["score"], item["profile_key"]))
    return scored, signals


_SUMMARY_SIGNAL_SUFFIX = "real-time services."

_SUMMARY_SIGNAL_TECH = [
    (["azure functions", "azure"], "Azure"),
    (["sql"], "SQL"),
    (["redis"], "Redis"),
]


def apply_jd_summary_signals(summary, keywords, archetype_name):
    """Append JD-matched tech signals to the summary ending for backend/swe2 archetypes.

    If the summary ends with 'real-time services.' and the JD mentions Azure, SQL,
    or Redis, the ending becomes 'real-time services using SQL, Redis, Azure.' etc.
    """
    if archetype_name not in {"backend"}:
        return summary
    stripped = (summary or "").rstrip()
    if not stripped.endswith(_SUMMARY_SIGNAL_SUFFIX):
        return summary
    kw_text = " ".join(normalize_text(k) for k in (keywords or []))
    matched = []
    for tokens, label in _SUMMARY_SIGNAL_TECH:
        if any(t in kw_text for t in tokens):
            matched.append(label)
    if not matched:
        return summary
    tech_str = ", ".join(matched)
    base = stripped[: -len(_SUMMARY_SIGNAL_SUFFIX)].rstrip()
    return f"{base} real-time services using {tech_str}."


def resolve_resume_profile(db_data, archetype_name, archetype, keywords=None, role_context=None):
    keywords = keywords or []
    role_context = role_context or {}
    default_profile = {
        "profile_key": default_profile_key_for_archetype(archetype_name),
        "subtitle": archetype["_meta"]["subtitle"],
        "summary": archetype["_meta"]["summary"],
    }
    rows = db_data.get("resume_profiles") or []
    if not rows:
        return default_profile
    if default_profile["profile_key"] not in build_resume_profile_index(db_data):
        return default_profile

    scored_profiles, signals = score_resume_profiles(db_data, archetype_name, keywords, role_context)
    best_non_blocked = next((item for item in scored_profiles if not item.get("blocked")), None)
    selected = best_non_blocked or next((item for item in scored_profiles if item["profile_key"] == default_profile["profile_key"]), None)
    if not selected:
        return default_profile

    profile_key = normalize_resume_profile_key(row_value(selected, "profile_key", default="")) or default_profile["profile_key"]
    subtitle = row_value(selected, "subtitle", default=default_profile["subtitle"])
    summary = row_value(selected, "summary", default=default_profile["summary"])
    summary = apply_jd_summary_signals(summary, keywords, archetype_name)
    fit_tier = "low_match_general" if signals.get("low_match_general_fallback") else ("adjacent_fit" if signals.get("fit_band_adjacent") else "direct_fit")
    return {
        "profile_key": profile_key,
        "subtitle": subtitle,
        "summary": summary,
        "fit_tier": fit_tier,
        "selection_reason": "profile_signal_rules",
        "profile_signals": signals,
        "profile_candidates": [
            {
                "profile_key": item["profile_key"],
                "score": item["score"],
                "blocked": item["blocked"],
                "matched_rules": item["matched_rules"],
            }
            for item in scored_profiles
        ],
    }


def build_resume_point_context(db_data):
    archetype_names = {
        row_value(row, "id"): row_value(row, "name", "label", default="")
        for row in db_data.get("archetypes", [])
        if row_value(row, "id") is not None
    }
    archetypes_by_point = {}
    for row in db_data.get("resume_point_archetypes", []):
        point_id = row_value(row, "resume_point_id", "point_id")
        if point_id is None:
            continue
        archetypes_by_point.setdefault(point_id, []).append(
            {
                "name": row_value(row, "archetype_name", "name", default="")
                or archetype_names.get(row_value(row, "archetype_id"), ""),
                "relevance": parse_int(row_value(row, "relevance", "score"), 0),
                "inclusion_mode": normalize_text(row_value(row, "inclusion_mode", default="context")) or "context",
            }
        )

    skills_by_point = {}
    for row in db_data.get("resume_point_skills", []):
        point_id = row_value(row, "resume_point_id", "point_id")
        if point_id is None:
            continue
        tokens = skills_by_point.setdefault(point_id, set())
        for value in [
            row_value(row, "skill_name"),
            row_value(row, "skill_normalized"),
            row_value(row, "capability_name"),
            row_value(row, "capability_normalized"),
            row_value(row, "name"),
            row_value(row, "normalized"),
        ]:
            normalized = normalize_text(value)
            if normalized:
                tokens.add(normalized)
    return archetypes_by_point, skills_by_point


def build_resume_point_variant_index(db_data):
    variants_by_point = {}
    for row in db_data.get("resume_point_variants", []):
        point_id = row_value(row, "resume_point_id", "point_id")
        if point_id is None:
            continue
        if not parse_bool(row_value(row, "approved", "is_approved"), True):
            continue
        variants_by_point.setdefault(point_id, []).append(dict(row))
    return variants_by_point


def build_variant_profile_preference_order(resume_profile):
    profile_key = normalize_resume_profile_key((resume_profile or {}).get("profile_key", ""))
    signals = (resume_profile or {}).get("profile_signals") or {}
    fit_tier = normalize_text((resume_profile or {}).get("fit_tier", ""))
    role_archetype_scores = signals.get("role_archetype_scores") or []
    role_archetype_map = build_role_archetype_score_map(role_archetype_scores)

    security_strength = parse_int(signals.get("jd_signal_security"), 0)
    backend_strength = parse_int(signals.get("jd_signal_backend"), 0)
    gameplay_strength = parse_int(signals.get("jd_signal_gameplay"), 0)
    liveops_strength = parse_int(signals.get("jd_signal_liveops"), 0)
    vr_strength = parse_int(signals.get("jd_signal_vr"), 0)

    ordered = [profile_key]
    ranked_role_profiles = [
        key
        for key, _score in sorted(role_archetype_map.items(), key=lambda item: (-item[1], item[0]))
        if key
    ]
    ordered.extend(ranked_role_profiles)
    if role_archetype_map.get("cyber", 0.0) >= max(
        role_archetype_map.get("backend", 0.0),
        role_archetype_map.get("gameplay", 0.0),
        role_archetype_map.get("liveops", 0.0),
        role_archetype_map.get("vr", 0.0),
        role_archetype_map.get("forensic", 0.0),
    ) or security_strength >= max(backend_strength, gameplay_strength, liveops_strength, vr_strength):
        ordered.extend(["forensic", "cyber", "backend", "general", "liveops", "vr", "gameplay"])
    elif role_archetype_map.get("backend", 0.0) >= max(
        role_archetype_map.get("gameplay", 0.0),
        role_archetype_map.get("liveops", 0.0),
        role_archetype_map.get("vr", 0.0),
    ) or backend_strength >= max(gameplay_strength, liveops_strength, vr_strength):
        ordered.extend(["backend", "liveops", "general", "cyber", "forensic", "gameplay", "vr"])
    elif role_archetype_map.get("liveops", 0.0) >= max(
        role_archetype_map.get("gameplay", 0.0),
        role_archetype_map.get("vr", 0.0),
    ) or liveops_strength >= max(gameplay_strength, vr_strength):
        ordered.extend(["liveops", "backend", "general", "gameplay", "cyber", "forensic", "vr"])
    elif role_archetype_map.get("vr", 0.0) > 0 or vr_strength > 0:
        ordered.extend(["vr", "general", "backend", "gameplay", "liveops", "cyber", "forensic"])
    else:
        ordered.extend(["gameplay", "liveops", "backend", "general", "vr", "cyber", "forensic"])

    if fit_tier in {"adjacent-fit", "adjacent_fit", "low-match-general", "low_match_general"}:
        ordered = [item for item in ordered if item != "gameplay"] + ["gameplay"]

    deduped = []
    seen = set()
    for item in ordered:
        normalized = normalize_resume_profile_key(item)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped


def generic_variant_penalty(text, resume_profile):
    normalized = normalize_text(text)
    fit_tier = normalize_text((resume_profile or {}).get("fit_tier", ""))
    if fit_tier not in {"adjacent-fit", "adjacent_fit", "low-match-general", "low_match_general"}:
        return 0

    gameplay_tokens = ["gameplay", "player-facing", "player facing", "custom game modes", "map vote", "cosmetics system", "rhythm-based", "immersive player", "multiplayer title"]
    penalty = 0
    if any(token in normalized for token in gameplay_tokens):
        penalty += 20
    if "gorilla tag" in normalized and "10m+ players" in normalized:
        penalty += 10
    return penalty


def resolve_resume_point_text(point_row, variants, archetype_name, archetype, resume_profile, variant_type="resume"):
    profile_key = normalize_resume_profile_key((resume_profile or {}).get("profile_key", ""))
    preferred_profiles = build_variant_profile_preference_order(resume_profile)
    candidate_names = {
        normalize_text(archetype_name),
        normalize_text(archetype.get("id", "")),
        normalize_text(archetype.get("label", "")),
        normalize_text(ARCHETYPE_ID_MAP.get(archetype_name, "")),
    }
    candidate_names.discard("")

    archetype_lookup = {
        row_value(row, "id"): normalize_text(row_value(row, "name", "label", default=""))
        for row in (archetype.get("_db_rows") or [])
        if row_value(row, "id") is not None
    }

    ranked = []
    for variant in variants:
        if normalize_text(row_value(variant, "variant_type", default="resume")) != normalize_text(variant_type):
            continue
        variant_profile = normalize_resume_profile_key(row_value(variant, "profile_key", default=""))
        variant_archetype_id = row_value(variant, "archetype_id")
        variant_archetype_name = archetype_lookup.get(variant_archetype_id, "")
        render_text = str(row_value(variant, "render_text", "text", default="")).strip()

        bucket = None
        profile_rank = preferred_profiles.index(variant_profile) if variant_profile in preferred_profiles else len(preferred_profiles)
        generic_penalty = generic_variant_penalty(render_text, resume_profile) if not variant_profile else 0
        if profile_key and variant_profile == profile_key and variant_archetype_name in candidate_names:
            bucket = 1
        elif profile_key and variant_profile == profile_key and variant_archetype_id is None:
            bucket = 2
        elif variant_profile and variant_profile in preferred_profiles and variant_archetype_id is None:
            bucket = 3
        elif not variant_profile and variant_archetype_name in candidate_names:
            bucket = 4
        elif parse_bool(row_value(variant, "is_default"), False):
            bucket = 5
        elif not variant_profile and variant_archetype_id is None:
            bucket = 6
        if bucket is None:
            continue

        ranked.append(
            (
                bucket,
                profile_rank,
                generic_penalty,
                -parse_int(row_value(variant, "is_default"), 0),
                parse_int(row_value(variant, "sort_priority"), 0),
                parse_int(row_value(variant, "id"), 0),
                render_text,
            )
        )

    if ranked:
        ranked.sort(key=lambda item: (item[0], item[1], item[2], item[3], item[4], item[5]))
        if ranked[0][6]:
            return ranked[0][6]

    for key in ["canonical_text", "text", "point_text", "body", "content"]:
        text = row_value(point_row, key, default="")
        if str(text).strip():
            return str(text).strip()
    return ""


def build_db_resume_points(db_data, archetype_name, archetype, resume_profile):
    experience_by_id = {}
    for row in db_data["experience"]:
        experience_by_id[row["id"]] = dict(row)

    projects = build_resume_project_index(db_data)
    archetypes_by_point, skills_by_point = build_resume_point_context(db_data)
    variants_by_point = build_resume_point_variant_index(db_data)

    points = []
    for row in db_data.get("resume_points", []):
        if not parse_bool(row_value(row, "approved", "is_approved"), True):
            continue
        experience_id = row_value(row, "experience_id", "exp_id")
        experience = experience_by_id.get(experience_id)
        if not experience:
            continue
        point_id = row_value(row, "id", "resume_point_id")
        text = resolve_resume_point_text(
            row,
            variants_by_point.get(point_id, []),
            archetype_name,
            archetype,
            resume_profile,
        )
        if not str(text).strip():
            continue
        project_id = row_value(row, "project_id")
        if project_id is not None and project_id not in projects:
            continue
        points.append(
            {
                "id": point_id,
                "experience_id": experience_id,
                "project_id": project_id,
                "text": str(text).strip(),
                "dedupe_key": row_value(row, "dedupe_key", default=normalize_text(text)),
                "importance": parse_int(row_value(row, "importance"), 0),
                "sort_priority": parse_int(row_value(row, "sort_priority", "priority", "display_order"), 0),
                "company": experience["company"],
                "title": experience["title"],
                "start_date": experience["start_date"],
                "end_date": experience["end_date"],
                "archetypes": archetypes_by_point.get(point_id, []),
                "skill_tokens": skills_by_point.get(point_id, set()),
            }
        )
    return points, projects


def keyword_match_score(point, keywords):
    if not keywords:
        return 0
    skill_tokens = point.get("skill_tokens") or set()
    text = normalize_text(point["text"])
    score = 0
    for keyword in keywords:
        if keyword in skill_tokens:
            score += 18
            continue
        if any(keyword in token or token in keyword for token in skill_tokens):
            score += 10
            continue
        if keyword in text:
            score += 6
    return score


def archetype_score(point, archetype_name, archetype):
    candidates = {
        normalize_text(archetype_name),
        normalize_text(archetype.get("id", "")),
        normalize_text(archetype.get("label", "")),
        normalize_text(ARCHETYPE_ID_MAP.get(archetype_name, "")),
    }
    candidates.discard("")
    entries = point.get("archetypes") or []
    matched = [entry for entry in entries if normalize_text(entry.get("name")) in candidates]
    selected = matched or [entry for entry in entries if entry.get("inclusion_mode") == "always"]
    if not selected:
        return 0, "context"
    best = max(selected, key=lambda entry: (entry.get("relevance", 0), entry.get("inclusion_mode", "")))
    return parse_int(best.get("relevance"), 0), normalize_text(best.get("inclusion_mode", "context")) or "context"


def weighted_archetype_score(point, archetype_name, archetype, role_archetype_scores=None):
    role_archetype_map = enrich_role_archetype_score_map(
        build_role_archetype_score_map(role_archetype_scores or []),
        archetype_name,
        archetype,
    )
    if not role_archetype_map:
        return archetype_score(point, archetype_name, archetype)

    entries = point.get("archetypes") or []
    weighted_total = 0.0
    weighted_entries = []
    always_entries = []
    for entry in entries:
        entry_key = normalize_archetype_score_key(entry.get("name"))
        if normalize_text(entry.get("inclusion_mode")) == "always":
            always_entries.append(entry)
        if not entry_key:
            continue
        role_weight = role_archetype_map.get(entry_key, 0.0)
        if role_weight <= 0:
            continue
        weighted_value = role_weight * parse_int(entry.get("relevance"), 0)
        weighted_total += weighted_value
        weighted_entries.append((weighted_value, entry))

    if weighted_entries:
        weighted_entries.sort(key=lambda item: (item[0], parse_int(item[1].get("relevance"), 0)), reverse=True)
        inclusion_mode = normalize_text(weighted_entries[0][1].get("inclusion_mode", "context")) or "context"
        return max(1, int(round(weighted_total))), inclusion_mode

    if always_entries:
        best = max(always_entries, key=lambda entry: parse_int(entry.get("relevance"), 0))
        return parse_int(best.get("relevance"), 0), normalize_text(best.get("inclusion_mode", "always")) or "always"

    return archetype_score(point, archetype_name, archetype)


def score_db_resume_point(point, archetype_name, archetype, keywords, role_archetype_scores=None):
    relevance, inclusion_mode = weighted_archetype_score(point, archetype_name, archetype, role_archetype_scores=role_archetype_scores)
    inclusion_weights = {
        "always": 1000,
        "prefer": 180,
        "context": 45,
        "hide": -100000,
    }
    skill_score = keyword_match_score(point, keywords)
    sort_priority = point.get("sort_priority", 0)
    total = (
        inclusion_weights.get(inclusion_mode, 0)
        + (relevance * 20)
        + (skill_score * 10)
        + (point.get("importance", 0) * 4)
        - sort_priority
    )
    return {
        **point,
        "relevance": relevance,
        "inclusion_mode": inclusion_mode,
        "skill_match_score": skill_score,
        "score": total,
    }


def build_experience_selection_budget(defaults, max_pages):
    base_role_budget = int(defaults.get("max_role_count", 4))
    base_per_project_budget = int(defaults.get("max_bullets_per_project", 3))
    base_project_budget = max(base_role_budget + 1, base_role_budget * 2)
    base_bullet_budget = max(base_project_budget * 2, base_role_budget * max(2, base_per_project_budget))

    if max_pages <= 1:
        page_pressure = "tight"
        role_budget = min(base_role_budget, 3)
        project_budget = min(base_project_budget, max(4, role_budget + 1))
        bullet_budget = base_bullet_budget  # no artificial cap; per-project limit governs
    else:
        page_pressure = "normal"
        role_budget = base_role_budget
        project_budget = base_project_budget
        bullet_budget = base_bullet_budget

    return {
        "page_pressure": page_pressure,
        "configured": {
            "roles": role_budget,
            "projects": project_budget,
            "bullets": bullet_budget,
            "max_bullets_per_project": base_per_project_budget,
        },
    }


def finalize_selected_experience_roles(role_pool, role_budget_details):
    def bullet_identity(bullet):
        explicit_id = str(bullet.get("id", "")).strip()
        if explicit_id:
            return explicit_id
        return normalize_text(bullet.get("text", ""))

    configured = role_budget_details["configured"]
    selected_role_pool = role_pool[: configured["roles"]]
    selected_project_ids: set[tuple[int, int]] = set()
    selected_projects_by_role: dict[int, list[dict[str, Any]]] = {index: [] for index in range(len(selected_role_pool))}
    remaining_project_budget = configured["projects"]

    for role_index, role in enumerate(selected_role_pool):
        if remaining_project_budget <= 0:
            break
        if role["direct_bullets"]:
            continue
        if not role["projects"]:
            continue
        selected_projects_by_role[role_index].append(role["projects"][0])
        selected_project_ids.add((role_index, role["projects"][0]["project_id"]))
        remaining_project_budget -= 1

    project_candidates = []
    for role_index, role in enumerate(selected_role_pool):
        for project_index, project in enumerate(role["projects"]):
            project_candidates.append((role_index, project_index, project))
    project_candidates.sort(
        key=lambda item: (
            -item[2]["score"],
            -item[2]["selected_point_count"],
            item[2]["sort_priority"],
            item[2]["label"].lower(),
        )
    )
    for role_index, _, project in project_candidates:
        if remaining_project_budget <= 0:
            break
        project_key = (role_index, project["project_id"])
        if project_key in selected_project_ids:
            continue
        selected_projects_by_role[role_index].append(project)
        selected_project_ids.add(project_key)
        remaining_project_budget -= 1

    bullet_slots: list[tuple[str, int, int | None, dict[str, Any], bool]] = []
    selected_bullet_ids: set[str] = set()
    selected_direct_bullets: dict[int, list[dict[str, Any]]] = {index: [] for index in range(len(selected_role_pool))}
    selected_project_bullets: dict[tuple[int, int], list[dict[str, Any]]] = {}

    for role_index, role in enumerate(selected_role_pool):
        chosen_projects = selected_projects_by_role[role_index]
        if chosen_projects:
            for project in chosen_projects:
                first_bullet = project["bullets"][0]
                selected_project_bullets[(role_index, project["project_id"])] = [first_bullet]
                selected_bullet_ids.add(bullet_identity(first_bullet))
        elif role["direct_bullets"]:
            first_bullet = role["direct_bullets"][0]
            selected_direct_bullets[role_index].append(first_bullet)
            selected_bullet_ids.add(bullet_identity(first_bullet))

    mandatory_bullet_count = len(selected_bullet_ids)
    bullet_budget = max(configured["bullets"], mandatory_bullet_count)

    for role_index, role in enumerate(selected_role_pool):
        for bullet in role["direct_bullets"]:
            bullet_id = bullet_identity(bullet)
            if bullet_id in selected_bullet_ids:
                continue
            bullet_slots.append(("direct", role_index, None, bullet, False))
        for project in selected_projects_by_role[role_index]:
            project_key = (role_index, project["project_id"])
            current_count = len(selected_project_bullets.get(project_key, []))
            for bullet in project["bullets"][current_count: configured["max_bullets_per_project"]]:
                bullet_id = bullet_identity(bullet)
                if bullet_id in selected_bullet_ids:
                    continue
                bullet_slots.append(("project", role_index, project["project_id"], bullet, bullet.get("inclusion_mode") == "prefer"))

    bullet_slots.sort(
        key=lambda item: (
            0 if item[4] else 1,
            -parse_int(item[3].get("score"), 0),
            -parse_int(item[3].get("relevance"), 0),
            -parse_int(item[3].get("importance"), 0),
            str(item[3].get("text", "")).lower(),
        )
    )

    for bucket_type, role_index, project_id, bullet, _ in bullet_slots:
        if len(selected_bullet_ids) >= bullet_budget:
            break
        bullet_id = bullet_identity(bullet)
        if bullet_id in selected_bullet_ids:
            continue
        if bucket_type == "project":
            project_key = (role_index, project_id)
            current = selected_project_bullets.setdefault(project_key, [])
            if len(current) >= configured["max_bullets_per_project"]:
                continue
            current.append(bullet)
        else:
            selected_direct_bullets[role_index].append(bullet)
        selected_bullet_ids.add(bullet_id)

    final_roles = []
    selected_project_count = 0
    for role_index, role in enumerate(selected_role_pool):
        direct_bullets = selected_direct_bullets.get(role_index, [])
        project_entries = []
        for project in selected_projects_by_role[role_index]:
            bullets = selected_project_bullets.get((role_index, project["project_id"]), [])
            if not bullets:
                continue
            project_entries.append({**project, "bullets": bullets})
        if not direct_bullets and not project_entries:
            continue
        selected_project_count += len(project_entries)
        final_roles.append(
            {
                "company": role["company"],
                "title": role["title"],
                "dates": role["dates"],
                "direct_bullets": direct_bullets,
                "projects": project_entries,
            }
        )

    role_budget_details["available"] = {
        "roles": len(role_pool),
        "projects": sum(len(role["projects"]) for role in role_pool),
        "bullets": sum(len(role["direct_bullets"]) for role in role_pool)
        + sum(len(project["bullets"]) for role in role_pool for project in role["projects"]),
    }
    role_budget_details["selected"] = {
        "roles": len(final_roles),
        "projects": selected_project_count,
        "bullets": len(selected_bullet_ids),
        "mandatory_bullets": mandatory_bullet_count,
    }
    return final_roles, role_budget_details


def build_selected_roles_from_resume_schema(db_data, archetype_name, archetype, keywords, max_pages, resume_profile=None):
    defaults = archetype["_meta"]["defaults"]
    max_bullets_per_project = int(defaults.get("max_bullets_per_project", 3))
    selection_budget = build_experience_selection_budget(defaults, max_pages)
    include_roles = archetype.get("experience", {}).get("include_roles", [])
    included_role_keys = {ROLE_ID_MAP[role_id] for role_id in include_roles if role_id in ROLE_ID_MAP}
    role_order = {ROLE_ID_MAP[role_id]: index for index, role_id in enumerate(include_roles) if role_id in ROLE_ID_MAP}

    keywords = [normalize_text(keyword) for keyword in keywords if normalize_text(keyword)]
    resume_profile = resume_profile or resolve_resume_profile(db_data, archetype_name, archetype, keywords=keywords)
    role_archetype_scores = ((resume_profile or {}).get("profile_signals") or {}).get("role_archetype_scores") or []
    archetype_with_db = dict(archetype)
    archetype_with_db["_db_rows"] = db_data.get("archetypes", [])
    resume_points, projects = build_db_resume_points(db_data, archetype_name, archetype_with_db, resume_profile)
    scored_points = []
    for point in resume_points:
        role_key = (point["company"], point["title"])
        if included_role_keys and role_key not in included_role_keys:
            continue
        scored = score_db_resume_point(
            point,
            archetype_name,
            archetype,
            keywords,
            role_archetype_scores=role_archetype_scores,
        )
        if scored["inclusion_mode"] == "hide":
            continue
        scored_points.append(scored)

    deduped = {}
    for point in scored_points:
        existing = deduped.get(point["dedupe_key"])
        if existing is None or (
            point["score"],
            point["relevance"],
            point["skill_match_score"],
            point["importance"],
            -point["sort_priority"],
            point["text"],
        ) > (
            existing["score"],
            existing["relevance"],
            existing["skill_match_score"],
            existing["importance"],
            -existing["sort_priority"],
            existing["text"],
        ):
            deduped[point["dedupe_key"]] = point

    grouped = {}
    for point in deduped.values():
        key = (point["company"], point["title"], point["start_date"], point["end_date"])
        entry = grouped.setdefault(
            key,
            {
                "company": point["company"],
                "title": point["title"],
                "dates": date_range(point["start_date"], point["end_date"]),
                "direct_bullets": [],
                "projects": {},
                "total_points": 0,
            },
        )
        entry["total_points"] += 1
        if point["project_id"] is None:
            entry["direct_bullets"].append(point)
            continue
        project = projects.get(point["project_id"])
        if not project:
            entry["direct_bullets"].append(point)
            continue
        project_entry = entry["projects"].setdefault(
            point["project_id"],
            {
                "project_id": point["project_id"],
                "label": project["label"],
                "stack": project["stack"],
                "sort_priority": project["sort_priority"],
                "bullets": [],
            },
        )
        project_entry["bullets"].append(point)

    role_pool = []
    for key, role in grouped.items():
        direct_bullets = sorted(
            role["direct_bullets"],
            key=lambda item: (-item["score"], -item["relevance"], -item["importance"], item["sort_priority"], item["text"].lower()),
        )[:max(1, max_bullets_per_project)]
        project_entries = []
        for project_entry in role["projects"].values():
            candidates = sorted(
                project_entry["bullets"],
                key=lambda item: (-item["score"], -item["relevance"], -item["importance"], item["sort_priority"], item["text"].lower()),
            )[:max(1, max_bullets_per_project)]
            if not candidates:
                continue
            project_entries.append(
                {
                    "project_id": project_entry["project_id"],
                    "label": project_entry["label"],
                    "stack": project_entry["stack"],
                    "sort_priority": project_entry["sort_priority"],
                    "bullets": candidates,
                    "selected_point_count": len(candidates),
                    "score": sum(item["score"] for item in candidates),
                }
            )

        project_entries.sort(
            key=lambda entry: (
                -entry["selected_point_count"],
                -entry["score"],
                entry["sort_priority"],
                entry["label"].lower(),
            )
        )
        if not project_entries and not direct_bullets:
            continue
        role_pool.append(
            {
                "company": role["company"],
                "title": role["title"],
                "dates": role["dates"],
                "direct_bullets": direct_bullets,
                "projects": project_entries,
                "_sort": (
                    role_order.get((role["company"], role["title"]), len(role_order)),
                    key[2] or "",
                    role["company"].lower(),
                ),
            }
        )

    role_pool.sort(key=lambda item: item["_sort"])
    for role in role_pool:
        role.pop("_sort", None)
    return finalize_selected_experience_roles(role_pool, selection_budget)


def build_selected_roles(db_data, archetype_name, archetype, keywords, max_pages, resume_profile=None):
    if has_resume_schema(db_data):
        return build_selected_roles_from_resume_schema(db_data, archetype_name, archetype, keywords, max_pages, resume_profile=resume_profile)

    defaults = archetype["_meta"]["defaults"]
    max_bullets_per_project = int(defaults.get("max_bullets_per_project", 3))
    selection_budget = build_experience_selection_budget(defaults, max_pages)
    role_lookup = role_lookup_by_id(db_data)

    keywords = [normalize_text(keyword) for keyword in keywords if normalize_text(keyword)]
    pinned_fragments = [normalize_text(PINNED_TEXT_MAP[item]) for item in archetype.get("experience", {}).get("pinned_bullets", []) if item in PINNED_TEXT_MAP]
    preferred_tags = set(archetype.get("experience", {}).get("prefer_tags", []))
    resume_points = build_resume_points(db_data)

    role_pool = []
    include_roles = archetype.get("experience", {}).get("include_roles", [])
    include_projects = archetype.get("experience", {}).get("include_projects", [])
    for role_id in include_roles:
        role_row = role_lookup.get(role_id)
        if not role_row:
            continue
        role_points = [
            score_resume_point(point, preferred_tags, pinned_fragments, keywords)
            for point in resume_points
            if point["role_id"] == role_id and point["project_id"] in include_projects
        ]
        deduped = {}
        for point in role_points:
            existing = deduped.get(point["dedupe_key"])
            if existing is None or (point["score"], point["strength"], point["text"]) > (existing["score"], existing["strength"], existing["text"]):
                deduped[point["dedupe_key"]] = point

        project_entries = []
        for project_id in include_projects:
            project = project_config(project_id)
            if project["role"] != role_id:
                continue
            candidates = [point for point in deduped.values() if point["project_id"] == project_id]
            if not candidates:
                continue
            candidates.sort(key=lambda item: (-int(item["pinned"]), -item["score"], -item["strength"], item["text"].lower()))
            selected_points = candidates[:max_bullets_per_project]
            project_entries.append(
                {
                    "project_id": project_id,
                    "label": project["label"],
                    "stack": project["stack"],
                    "sort_priority": len(project_entries),
                    "bullets": selected_points,
                    "selected_point_count": len(selected_points),
                    "score": sum(item["score"] for item in selected_points),
                }
            )
        if not project_entries:
            continue
        project_entries.sort(key=lambda entry: (-entry["selected_point_count"], -entry["score"], entry["label"].lower()))
        role_pool.append(
            {
                "company": role_row["company"],
                "title": role_row["title"],
                "dates": date_range(role_row["start_date"], role_row["end_date"]),
                "direct_bullets": [],
                "projects": project_entries,
            }
        )
    return finalize_selected_experience_roles(role_pool, selection_budget)


def build_education_entries(db_data, archetype, keywords, runtime_feature_map=None):
    include_ids = archetype.get("education", {}).get("include_ids", [])
    allowed = {EDUCATION_ID_MAP[item] for item in include_ids if item in EDUCATION_ID_MAP}
    runtime_feature_map = runtime_feature_map or {}
    cerritos_allowed = parse_int(runtime_feature_map.get("cerritos_fit"), 0) > 0 and (
        parse_int(runtime_feature_map.get("cluster_scores", {}).get("networking"), 0) > 0
        or parse_int(runtime_feature_map.get("cluster_scores", {}).get("security"), 0) > 0
        or parse_int(runtime_feature_map.get("cluster_scores", {}).get("forensics"), 0) > 0
    )
    usc_rows = []
    entries = []
    for row in db_data["education"]:
        key = f"{row['school']}|{row['degree']}|{row['field']}"
        if _PRIMARY_SCHOOL and row["school"] == _PRIMARY_SCHOOL:
            usc_rows.append(dict(row))
            continue
        if allowed and key not in allowed:
            continue
        if _SUPPLEMENTAL_SCHOOL and row["school"] == _SUPPLEMENTAL_SCHOOL and not cerritos_allowed:
            continue
        detail = f"{row['degree']} {row['field']}".strip()
        if row.get("status") and row["status"] != "completed":
            detail = f"{detail} ({row['status'].replace('-', ' ')})"
        entries.append({"school": row["school"], "detail": detail, "year": ""})
    usc_entry = build_usc_education_entry(usc_rows, allowed, runtime_feature_map=runtime_feature_map)
    if usc_entry:
        entries.insert(0, usc_entry)
    return entries


def games_field_role_fit(runtime_feature_map):
    cluster_scores = (runtime_feature_map or {}).get("cluster_scores", {}) or {}
    gameplay = parse_int(cluster_scores.get("gameplay"), 0)
    liveops = parse_int(cluster_scores.get("liveops"), 0)
    xr = parse_int(cluster_scores.get("xr"), 0)
    backend = parse_int(cluster_scores.get("backend"), 0)
    frontend = parse_int(cluster_scores.get("frontend"), 0)
    systems = parse_int(cluster_scores.get("systems"), 0)

    game_total = gameplay + liveops + xr
    non_game_total = backend + frontend + systems
    direct_match_density = float((runtime_feature_map or {}).get("direct_match_density") or 0.0)

    if game_total >= 4:
        return True
    if gameplay >= 2 or xr >= 2:
        return True
    if game_total > non_game_total and direct_match_density >= 0.15:
        return True
    return False


def select_usc_field_variant(row, runtime_feature_map=None):
    original_field = str(row.get("field") or "").strip()
    if normalize_text(original_field) != "computer science (games)":
        return original_field
    if games_field_role_fit(runtime_feature_map or {}):
        return "Computer Science (Games)"
    return "Computer Science"


def build_usc_education_entry(usc_rows, allowed, runtime_feature_map=None):
    if not usc_rows:
        return None

    primary_allowed_keys = _PRIMARY_SCHOOL_ALLOWED_KEYS
    if allowed and primary_allowed_keys and not allowed.intersection(primary_allowed_keys):
        return None

    preferred_field = "Computer Science (Games)" if games_field_role_fit(runtime_feature_map or {}) else "Computer Science"
    preferred_row = None
    fallback_row = None
    for row in usc_rows:
        if normalize_text(row.get("field")) == normalize_text(preferred_field):
            preferred_row = dict(row)
            break
        if fallback_row is None and normalize_text(row.get("field")) in {"computer science", "computer science (games)"}:
            fallback_row = dict(row)

    selected_row = preferred_row or fallback_row or dict(usc_rows[0])
    detail = f"{selected_row.get('degree', '').strip()} {preferred_field}".strip()
    if selected_row.get("status") and selected_row["status"] != "completed":
        detail = f"{detail} ({selected_row['status'].replace('-', ' ')})"
    return {
        "school": selected_row.get("school", _PRIMARY_SCHOOL),
        "detail": detail,
        "year": "",
    }


def explain_education_selection(db_data, selected_entries, runtime_feature_map=None):
    runtime_feature_map = runtime_feature_map or {}
    selected_schools = {entry.get("school") for entry in selected_entries}
    cerritos_selected = bool(_SUPPLEMENTAL_SCHOOL) and _SUPPLEMENTAL_SCHOOL in selected_schools
    cerritos_fit = parse_int(runtime_feature_map.get("cerritos_fit"), 0)
    cluster_scores = runtime_feature_map.get("cluster_scores", {}) or {}
    cerritos_signal_support = {
        "networking": parse_int(cluster_scores.get("networking"), 0),
        "security": parse_int(cluster_scores.get("security"), 0),
        "forensics": parse_int(cluster_scores.get("forensics"), 0),
    }
    return {
        "selected": list(selected_entries),
        "cerritos": {
            "included": cerritos_selected,
            "fit_score": cerritos_fit,
            "supporting_cluster_scores": cerritos_signal_support,
            "reason": (
                "Included because networking/security/forensics signals made the additional education relevant."
                if cerritos_selected
                else "Excluded because role signals did not justify extra education space."
            ),
        },
    }


def certification_role_relevance(certification, runtime_feature_map):
    cluster_scores = runtime_feature_map.get("cluster_scores", {})
    cert_name = normalize_text(certification.get("name"))
    cert_detail = normalize_text(certification.get("issuer"))
    score = 0

    if "ccna" in cert_name:
        score += 10
    if any(token in cert_name or token in cert_detail for token in ["network", "routing", "switching", "wireless", "cisco"]):
        score += parse_int(cluster_scores.get("networking"), 0) * 2
    if any(token in cert_name or token in cert_detail for token in ["cyber", "security"]):
        score += parse_int(cluster_scores.get("security"), 0) * 2
    if any(token in cert_name or token in cert_detail for token in ["forensic", "incident"]):
        score += parse_int(cluster_scores.get("forensics"), 0) * 2

    if normalize_text(certification.get("status")) == "earned":
        score += 2
    return score


def build_certification_entries(db_data, archetype, keywords, runtime_feature_map=None):
    include_ids = archetype.get("certifications", {}).get("include_ids", [])
    allowed = {CERT_ID_MAP[item] for item in include_ids if item in CERT_ID_MAP}
    runtime_feature_map = runtime_feature_map or {}
    cert_requested = bool(runtime_feature_map.get("cert_requested"))
    matched_cert_requirement_count = parse_int(runtime_feature_map.get("matched_cert_requirement_count"), 0)
    if not cert_requested or matched_cert_requirement_count <= 0:
        return []

    max_items = 1
    if parse_int(runtime_feature_map.get("cert_sensitivity"), 0) >= 12 or matched_cert_requirement_count >= 2:
        max_items = 2

    ranked = []
    for row in db_data["certifications"]:
        if allowed and row["name"] not in allowed:
            continue
        relevance = certification_role_relevance(row, runtime_feature_map)
        if relevance <= 0:
            continue
        suffix = ""
        if row.get("status") and row["status"] != "earned":
            suffix = f" | {row['status'].replace('-', ' ')}"
        ranked.append(
            {
                "name": row["name"],
                "detail": f"{row['issuer']}{suffix}",
                "year": month_label(row.get("date")) if row.get("date") else "",
                "_relevance": relevance,
            }
        )
    ranked.sort(key=lambda item: (-item["_relevance"], item["name"].lower()))
    return [{key: value for key, value in item.items() if key != "_relevance"} for item in ranked[:max_items]]


def explain_certification_selection(db_data, selected_entries, runtime_feature_map=None):
    runtime_feature_map = runtime_feature_map or {}
    selected_names = {entry.get("name") for entry in selected_entries}
    cert_requested = bool(runtime_feature_map.get("cert_requested"))
    matched_cert_requirement_count = parse_int(runtime_feature_map.get("matched_cert_requirement_count"), 0)
    cert_sensitivity = parse_int(runtime_feature_map.get("cert_sensitivity"), 0)
    decisions = []
    for row in db_data.get("certifications", []):
        relevance = certification_role_relevance(row, runtime_feature_map)
        included = row.get("name") in selected_names
        decisions.append(
            {
                "name": row.get("name"),
                "included": included,
                "relevance": relevance,
                "reason": (
                    "Included because certification request/sensitivity justified resume space."
                    if included
                    else "Excluded because certification relevance or request strength was too low."
                ),
            }
        )
    return {
        "requested": cert_requested,
        "matched_requirement_count": matched_cert_requirement_count,
        "cert_sensitivity": cert_sensitivity,
        "selected": list(selected_entries),
        "decisions": decisions,
    }


def replace_placeholders(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)


def get_section_indices(paragraphs):
    sections = {}
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip().upper()
        if text in {"SKILLS", "EXPERIENCE", "EDUCATION", "CERTIFICATIONS"}:
            sections[text] = index
    return sections


SECTION_LABEL_BY_ID = {
    "skills": "SKILLS",
    "experience": "EXPERIENCE",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
}


def paragraph_has_numbering(paragraph):
    try:
        return paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    except Exception:
        return False


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def clone_paragraph_before(anchor, donor):
    new_p = deepcopy(donor._p)
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    strip_paragraph_section_break(paragraph)
    return paragraph


def clear_runs(paragraph):
    children = list(paragraph._p)
    ppr = paragraph._p.pPr
    for child in children:
        if ppr is not None and child == ppr:
            continue
        paragraph._p.remove(child)


def copy_run_style(target_run, donor_run):
    target_run.bold = donor_run.bold
    target_run.italic = donor_run.italic
    target_run.underline = donor_run.underline
    if donor_run.font.name:
        target_run.font.name = donor_run.font.name
    if donor_run.font.size:
        target_run.font.size = donor_run.font.size
    if donor_run.font.color and donor_run.font.color.rgb:
        target_run.font.color.rgb = donor_run.font.color.rgb
    donor_rpr = donor_run._r.rPr
    if donor_rpr is not None:
        if target_run._r.rPr is not None:
            target_run._r.remove(target_run._r.rPr)
        target_run._r.insert(0, deepcopy(donor_rpr))


def strip_paragraph_section_break(paragraph):
    ppr = paragraph._p.pPr
    if ppr is None:
        return
    sect_pr = ppr.sectPr
    if sect_pr is not None:
        ppr.remove(sect_pr)


def parse_bold_segments(text: str) -> list[tuple[str, int, bool]]:
    """Split text on **bold** markers into (text, donor_index=0, is_bold) tuples."""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    result = []
    for i, part in enumerate(parts):
        if part:
            result.append((part, 0, i % 2 == 1))
    return result if result else [(text, 0, False)]


def set_runs(paragraph, segments, donor_runs):
    clear_runs(paragraph)
    if not donor_runs:
        for item in segments:
            text = item[0]
            paragraph.add_run(text)
        return
    for item in segments:
        text = item[0]
        donor_index = item[1]
        is_bold = item[2] if len(item) > 2 else False
        if not text:
            continue
        run = paragraph.add_run(text)
        source = donor_runs[min(donor_index, len(donor_runs) - 1)]
        copy_run_style(run, source)
        if is_bold:
            run.bold = True


def find_donors(document):
    paragraphs = document.paragraphs
    sections = get_section_indices(paragraphs)
    if "SKILLS" not in sections or "EXPERIENCE" not in sections or "EDUCATION" not in sections:
        fail("Template is missing required headings: SKILLS / EXPERIENCE / EDUCATION")
    skills_idx = sections["SKILLS"]
    experience_idx = sections["EXPERIENCE"]
    education_idx = sections["EDUCATION"]
    cert_idx = sections.get("CERTIFICATIONS")

    skills_body = paragraphs[skills_idx + 1 : experience_idx]
    experience_body = paragraphs[experience_idx + 1 : education_idx]
    education_body = paragraphs[education_idx + 1 : cert_idx if cert_idx is not None else len(paragraphs)]

    skill_row = next((p for p in skills_body if ":" in p.text.strip()), None)
    role_header = next((p for p in experience_body if "|" in p.text and not paragraph_has_numbering(p)), None)
    role_header_index = experience_body.index(role_header) if role_header else -1
    date_row = None
    project_row = None
    bullet_row = next((p for p in experience_body if paragraph_has_numbering(p)), None)
    if role_header_index >= 0:
        for p in experience_body[role_header_index + 1 :]:
            text = p.text.strip()
            if not text or paragraph_has_numbering(p) or "|" in text:
                continue
            if project_row is None and not re.search(r"\b\d{4}\b", text):
                project_row = p
                continue
            if date_row is None:
                date_row = p
                if project_row is not None:
                    break
    education_row = next((p for p in education_body if p.text.strip()), None)
    if not skill_row or not role_header or not bullet_row or not education_row:
        fail("Template structure does not match expected resume layout.")
    return {
        "sections": sections,
        "skill_row": skill_row,
        "role_header": role_header,
        "date_row": date_row,
        "project_row": project_row,
        "bullet_row": bullet_row,
        "education_row": education_row,
        "heading_row": paragraphs[education_idx],
    }


def replace_section(document, start_label, end_label, builder):
    paragraphs = document.paragraphs
    sections = get_section_indices(paragraphs)
    start_index = sections[start_label]
    has_end = end_label in sections
    end_index = sections[end_label] if has_end else len(paragraphs)
    body = paragraphs[start_index + 1 : end_index]
    if not body and not has_end:
        fail(f"Section {start_label} has no body to rewrite.")
    if has_end:
        anchor = paragraphs[end_index]
        for paragraph in reversed(body):
            remove_paragraph(paragraph)
    else:
        anchor = body[-1]
        for paragraph in reversed(body[:-1]):
            remove_paragraph(paragraph)
    builder(anchor)
    if not has_end:
        remove_paragraph(anchor)


def apply_section_plan_to_document(document, resume_plan):
    if not resume_plan:
        return
    final_sections = getattr(resume_plan, "final_sections", None) or []
    if not final_sections:
        return
    sections = get_section_indices(document.paragraphs)
    current_labels = [label for label, _ in sorted(sections.items(), key=lambda item: item[1])]
    if not current_labels:
        return

    visible_labels = []
    for section in final_sections:
        if not getattr(section, "visible", False):
            continue
        label = SECTION_LABEL_BY_ID.get(getattr(section, "section_id", ""))
        if label and label in sections and label not in visible_labels:
            visible_labels.append(label)

    if not visible_labels:
        return

    paragraphs = document.paragraphs
    sections = get_section_indices(paragraphs)
    ordered_current = [label for label, _ in sorted(sections.items(), key=lambda item: item[1])]
    first_index = sections[ordered_current[0]]
    last_end = first_index
    block_ranges = {}
    for index, label in enumerate(ordered_current):
        start_index = sections[label]
        next_label = ordered_current[index + 1] if index + 1 < len(ordered_current) else None
        end_index = sections[next_label] if next_label else len(paragraphs)
        block_ranges[label] = (start_index, end_index)
        last_end = max(last_end, end_index)

    parent = paragraphs[first_index]._element.getparent()
    tail_anchor = paragraphs[last_end]._element if last_end < len(paragraphs) else None
    blocks = {}
    for label in ordered_current:
        start_index, end_index = block_ranges[label]
        block_elements = [paragraphs[item_index]._element for item_index in range(start_index, end_index)]
        blocks[label] = block_elements

    for label in ordered_current:
        if label not in visible_labels:
            for element in blocks[label]:
                parent.remove(element)

    insert_labels = [label for label in visible_labels if label in blocks]
    insert_index = parent.index(tail_anchor) if tail_anchor is not None else len(parent)
    for label in ordered_current:
        if label in insert_labels:
            for element in blocks[label]:
                parent.remove(element)
    for label in insert_labels:
        for element in blocks[label]:
            parent.insert(insert_index, element)
            insert_index += 1


SKILL_GROUP_CHAR_LIMIT = 78


def format_skill_group_items(items):
    primary = [f"{item['display']} (primary)" for item in items if item.get("emphasis") == "primary" and item.get("display")]
    experience = [item["display"] for item in items if item.get("emphasis") == "experience" and item.get("display")]
    plain = [item["display"] for item in items if item.get("emphasis") not in {"primary", "experience"} and item.get("display")]

    rendered = []
    rendered.extend(primary)
    if experience:
        rendered.append(f"experience with {', '.join(experience)}")
    rendered.extend(plain)
    return ", ".join(rendered)


def trim_skill_group_items(entry, char_limit=SKILL_GROUP_CHAR_LIMIT):
    items = list(entry.get("items") or [])
    if not items:
        return []

    label = f"{entry.get('label', '')}: "
    while len(items) > 1:
        formatted = format_skill_group_items(items)
        if len(label) + len(formatted) <= char_limit:
            break
        items.pop()
    return items


def build_skills_section(document, donors, entries):
    def builder(anchor):
        rendered = []
        for entry in entries:
            items = trim_skill_group_items(entry)
            if not items:
                continue
            paragraph = clone_paragraph_before(anchor, donors["skill_row"])
            donor_runs = paragraph.runs[:] or donors["skill_row"].runs[:]
            set_runs(
                paragraph,
                [(f"{entry['label']}: ", 0), (format_skill_group_items(items), 1 if len(donor_runs) > 1 else 0)],
                donor_runs,
            )
            rendered.append(paragraph)

        if rendered:
            rendered[-1].paragraph_format.space_after = Pt(8)

    replace_section(document, "SKILLS", "EXPERIENCE", builder)


def build_experience_section(document, donors, roles, is_games=True):
    def ensure_right_tab_stop(paragraph):
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        paragraph.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    def set_keep_with_next(paragraph, value=True):
        paragraph.paragraph_format.keep_with_next = value

    def set_keep_together(paragraph, value=True):
        paragraph.paragraph_format.keep_together = value

    def builder(anchor):
        for role in roles:
            header = clone_paragraph_before(anchor, donors["role_header"])
            donor_runs = header.runs[:] or donors["role_header"].runs[:]
            clear_runs(header)
            company_run = header.add_run(f"{role['company']} ")
            copy_run_style(company_run, donor_runs[0])

            title_run = header.add_run(f"| {role['title']}")
            copy_run_style(title_run, donor_runs[1] if len(donor_runs) > 1 else donor_runs[0])

            if role.get("dates"):
                ensure_right_tab_stop(header)
                tab_run = header.add_run("\t")
                copy_run_style(tab_run, donor_runs[0])
                date_run = header.add_run(role["dates"])
                date_donor_runs = (donors["date_row"].runs[:] if donors.get("date_row") is not None else []) or donor_runs
                copy_run_style(date_run, date_donor_runs[0])
            if role["projects"] or role.get("direct_bullets"):
                set_keep_with_next(header, True)
            if is_games:
                for project in role["projects"]:
                    project_donor = donors["project_row"] or donors["bullet_row"]
                    project_paragraph = clone_paragraph_before(anchor, project_donor)
                    project_runs = project_paragraph.runs[:] or project_donor.runs[:]
                    set_runs(
                        project_paragraph,
                        [(f"  {project['label']}", 0), (f" ({project['stack']})", 1 if len(project_runs) > 1 else 0)],
                        project_runs,
                    )
                    set_keep_with_next(project_paragraph, True)
                    for index, bullet in enumerate(project["bullets"]):
                        bullet_paragraph = clone_paragraph_before(anchor, donors["bullet_row"])
                        set_runs(bullet_paragraph, parse_bold_segments(bullet["text"]), bullet_paragraph.runs[:] or donors["bullet_row"].runs[:])
                        set_keep_together(bullet_paragraph, True)
                        is_last_project_bullet = index == len(project["bullets"]) - 1
                        set_keep_with_next(bullet_paragraph, not is_last_project_bullet)
                for bullet in role.get("direct_bullets", []):
                    bullet_paragraph = clone_paragraph_before(anchor, donors["bullet_row"])
                    set_runs(bullet_paragraph, parse_bold_segments(bullet["text"]), bullet_paragraph.runs[:] or donors["bullet_row"].runs[:])
                    set_keep_together(bullet_paragraph, True)
            else:
                # Non-game format: flatten all project bullets + direct bullets into a single list
                all_bullets = [b for project in role["projects"] for b in project["bullets"]]
                all_bullets += list(role.get("direct_bullets", []))
                for bullet in all_bullets:
                    bullet_paragraph = clone_paragraph_before(anchor, donors["bullet_row"])
                    set_runs(bullet_paragraph, parse_bold_segments(bullet["text"]), bullet_paragraph.runs[:] or donors["bullet_row"].runs[:])
                    set_keep_together(bullet_paragraph, True)

    replace_section(document, "EXPERIENCE", "EDUCATION", builder)


def build_education_section(document, donors, entries):
    def builder(anchor):
        for entry in entries:
            paragraph = clone_paragraph_before(anchor, donors["education_row"])
            donor_runs = paragraph.runs[:] or donors["education_row"].runs[:]
            set_runs(
                paragraph,
                [
                    (f"{entry['school']} - ", 0),
                    (entry["detail"], 1 if len(donor_runs) > 1 else 0),
                    ((f"    {entry['year']}") if entry["year"] else "", 1 if len(donor_runs) > 1 else 0),
                ],
                donor_runs,
            )

    end_label = "CERTIFICATIONS" if "CERTIFICATIONS" in get_section_indices(document.paragraphs) else "__END__"
    replace_section(document, "EDUCATION", end_label, builder)


def append_certifications_section(document, donors, entries):
    if not entries:
        return
    sections = get_section_indices(document.paragraphs)
    if "CERTIFICATIONS" in sections:
        return
    anchor = document.paragraphs[-1]
    heading = clone_paragraph_before(anchor, donors["heading_row"])
    set_runs(heading, [("CERTIFICATIONS", 0)], heading.runs[:] or donors["heading_row"].runs[:])
    for entry in entries:
        paragraph = clone_paragraph_before(anchor, donors["education_row"])
        donor_runs = paragraph.runs[:] or donors["education_row"].runs[:]
        set_runs(
            paragraph,
            [
                (f"{entry['name']} - ", 0),
                (entry["detail"], 1 if len(donor_runs) > 1 else 0),
                ((f"    {entry['year']}") if entry["year"] else "", 1 if len(donor_runs) > 1 else 0),
            ],
            donor_runs,
        )


def apply_page_format(document, page_format):
    if page_format == "letter":
        width, height = Inches(8.5), Inches(11)
    else:
        width, height = Inches(8.27), Inches(11.69)
    for section in document.sections:
        section.page_width = width
        section.page_height = height


def normalize_section_columns(document):
    for paragraph in document.paragraphs:
        ppr = paragraph._p.pPr
        if ppr is None:
            continue
        sect_pr = ppr.sectPr
        if sect_pr is None:
            continue
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            sect_pr.remove(cols)
    for section in document.sections:
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            sect_pr.remove(cols)


def make_json_safe(value):
    if hasattr(value, "__dataclass_fields__"):
        return make_json_safe(asdict(value))
    if isinstance(value, dict):
        return {str(key): make_json_safe(item) for key, item in value.items()}
    if isinstance(value, set):
        return sorted(make_json_safe(item) for item in value)
    if isinstance(value, (list, tuple)):
        return [make_json_safe(item) for item in value]
    return value


def estimate_text_lines(text, width=80):
    normalized = str(text or "").strip()
    if not normalized:
        return 0
    return max(1, (len(normalized) // width) + (1 if len(normalized) % width else 0))


def allowed_lines_for_plan(page_format, max_pages):
    per_page = 64 if page_format == "letter" else 68
    return per_page * max(1, int(max_pages or 1))


def pressure_bucket(estimated_lines, allowed_lines):
    if allowed_lines <= 0:
        return "critical"
    ratio = estimated_lines / allowed_lines
    if ratio <= 0.82:
        return "low"
    if ratio <= 0.95:
        return "medium"
    if ratio <= 1.05:
        return "high"
    return "critical"


def sufficiency_bucket(count):
    if count <= 0:
        return "missing"
    if count == 1:
        return "weak"
    if count <= 3:
        return "adequate"
    return "strong"


def fallback_layout_rows():
    return [
        {
            "layout_key": "standard-1p",
            "layout_family": LAYOUT_GROUPED_STANDARD,
            "label": "Standard One Page",
            "page_count": 1,
            "max_roles": 4,
            "max_projects": 7,
            "max_bullets_total": 20,
            "max_bullets_per_role": 8,
            "max_bullets_per_project": 6,
            "max_direct_bullets_per_role": 1,
            "skill_group_char_limit": 78,
            "max_education_items": 2,
            "max_certification_items": 2,
            "certification_policy": "contextual",
            "section_order": "summary,skills,experience,education,certifications",
            "min_role_score": 0,
            "min_project_score": 0,
            "min_bullet_score": 0,
            "allow_low_match_context": 0,
            "approved": 1,
        },
        {
            "layout_key": "standard-2p",
            "layout_family": LAYOUT_GROUPED_STANDARD,
            "label": "Standard Two Page",
            "page_count": 2,
            "max_roles": 4,
            "max_projects": 8,
            "max_bullets_total": 30,
            "max_bullets_per_role": 10,
            "max_bullets_per_project": 8,
            "max_direct_bullets_per_role": 2,
            "skill_group_char_limit": 90,
            "max_education_items": 2,
            "max_certification_items": 3,
            "certification_policy": "contextual",
            "section_order": "summary,skills,experience,education,certifications",
            "min_role_score": 0,
            "min_project_score": 0,
            "min_bullet_score": 0,
            "allow_low_match_context": 0,
            "approved": 1,
        },
    ]


def parse_layout_section_order(layout_row):
    raw = str(layout_row.get("section_order") or "").strip()
    if not raw:
        return list(LAYOUT_STABLE_SECTION_ORDER)
    seen = set()
    ordered = []
    for part in raw.split(","):
        section_id = normalize_text(part)
        if section_id not in {"summary", "skills", "experience", "education", "certifications"}:
            continue
        if section_id in seen:
            continue
        seen.add(section_id)
        ordered.append(section_id)
    for section_id in LAYOUT_STABLE_SECTION_ORDER:
        if section_id not in seen:
            ordered.append(section_id)
    return ordered


def build_layout_profile_rule_index(db_data, profile_key):
    normalized_profile = normalize_resume_profile_key(profile_key)
    allowed = []
    for row in db_data.get("resume_layout_profile_rules", []):
        if not parse_bool(row.get("approved"), True):
            continue
        if normalize_resume_profile_key(row.get("profile_key")) != normalized_profile:
            continue
        allowed.append(dict(row))
    return allowed


def build_layout_signal_rule_index(db_data):
    rules = {}
    for row in db_data.get("resume_layout_signal_rules", []):
        if not parse_bool(row.get("approved"), True):
            continue
        layout_key = row.get("layout_key")
        if not layout_key:
            continue
        rules.setdefault(layout_key, []).append(dict(row))
    return rules


def build_layout_signal_map(candidate_content, runtime_feature_map):
    strong_project_count = 0
    selected_project_count = 0
    weak_context_bullet_count = 0
    for role in candidate_content.roles:
        for project in role.projects:
            selected_project_count += 1
            if project.score >= 20:
                strong_project_count += 1
            for bullet in project.bullets:
                if bullet.inclusion_mode == "context" and bullet.score < 80:
                    weak_context_bullet_count += 1
        for bullet in role.direct_bullets:
            if bullet.inclusion_mode == "context" and bullet.score < 80:
                weak_context_bullet_count += 1

    requirement_counts = (runtime_feature_map or {}).get("requirement_counts") or {}
    role_archetype_scores = (runtime_feature_map or {}).get("role_archetype_scores") or []
    role_archetype_map = build_role_archetype_score_map(role_archetype_scores)
    signals = {
        "direct_keyword_match_density": float((runtime_feature_map or {}).get("direct_match_density") or 0.0),
        "requirement_count_matched": parse_int(requirement_counts.get("matched"), 0),
        "requirement_count_unmatched": parse_int(requirement_counts.get("unmatched"), 0),
        "matched_cert_requirement_count": parse_int((runtime_feature_map or {}).get("matched_cert_requirement_count"), 0),
        "selected_project_count_candidate": selected_project_count,
        "strong_project_count_candidate": strong_project_count,
        "weak_context_bullet_count_candidate": weak_context_bullet_count,
        "cert_sensitivity": parse_int((runtime_feature_map or {}).get("cert_sensitivity"), 0),
        "education_sensitivity": parse_int((runtime_feature_map or {}).get("education_sensitivity"), 0),
    }
    for archetype_key, score in role_archetype_map.items():
        signals[f"role_archetype_score_{archetype_key}"] = round(float(score), 4)
    if role_archetype_scores:
        signals[f"role_primary_archetype_is_{role_archetype_scores[0]['archetype_key']}"] = 1
    return signals


def build_layout_section_caps(layout_row, candidate_content):
    skill_cap = 5
    if parse_int(layout_row.get("skill_group_char_limit"), 0) and parse_int(layout_row.get("skill_group_char_limit"), 0) < 70:
        skill_cap = min(skill_cap, 4)
    return {
        "summary": 1,
        "skills": min(skill_cap, max(1, len(candidate_content.skill_groups))),
        "experience": parse_int(layout_row.get("max_roles"), len(candidate_content.roles)),
        "education": min(parse_int(layout_row.get("max_education_items"), 2), max(1, len(candidate_content.education))) if candidate_content.education else 0,
        "certifications": min(parse_int(layout_row.get("max_certification_items"), 4), len(candidate_content.certifications)),
    }


def build_candidate_content(summary_text, skill_groups, roles, education_entries, certification_entries):
    summary_candidates = []
    if summary_text:
        summary_candidates.append(
            SummaryCandidate(
                id="summary.primary",
                text=summary_text,
                score=100,
                priority=100,
                min_keep=True,
                drop_cost=100,
                estimated_lines=estimate_text_lines(summary_text, 88),
                layout_affinity=[LAYOUT_GROUPED_STANDARD, LAYOUT_GROUPED_TOP_CERTS],
                selection_reason=["resume_profile_summary"],
                source_ids=["resume_profile"],
            )
        )

    skill_group_candidates = []
    for index, group in enumerate(skill_groups):
        item_candidates = []
        for item_index, item in enumerate(group.get("items", [])):
            item_candidates.append(
                SkillItemCandidate(
                    id=f"skill-item.{group.get('group_id')}.{item_index}",
                    display=item.get("display", item.get("name", "")),
                    score=parse_int(item.get("score"), 0),
                    priority=parse_int(item.get("item_rank"), 0),
                    min_keep=item_index == 0,
                    drop_cost=6 if item.get("matched_keyword") else 2,
                    estimated_lines=1,
                    selection_reason=["skill_group_member"],
                    source_ids=[item.get("name", "")],
                )
            )
        skill_group_candidates.append(
            SkillGroupCandidate(
                id=f"skill-group.{group.get('group_id')}",
                label=group.get("label", group.get("group_id", "").title()),
                score=parse_int(group.get("score"), 0),
                priority=100 - index,
                min_keep=index == 0,
                drop_cost=10,
                estimated_lines=1,
                layout_affinity=[LAYOUT_GROUPED_STANDARD, LAYOUT_GROUPED_TOP_CERTS],
                selection_reason=["selected_skill_group"],
                source_ids=[group.get("group_id", "")],
                items=item_candidates,
            )
        )

    role_candidates = []
    for role_index, role in enumerate(roles):
        direct_bullets = []
        for bullet_index, bullet in enumerate(role.get("direct_bullets", [])):
            direct_bullets.append(
                BulletCandidate(
                    id=f"bullet.role.{role_index}.direct.{bullet_index}",
                    text=bullet.get("text", ""),
                    score=parse_int(bullet.get("score"), 0),
                    priority=parse_int(bullet.get("importance"), 0),
                    min_keep=bullet_index == 0,
                    drop_cost=8,
                    estimated_lines=estimate_text_lines(bullet.get("text", ""), 96),
                    inclusion_mode=normalize_text(bullet.get("inclusion_mode", "context")) or "context",
                    selection_reason=["selected_direct_bullet"],
                    source_ids=[str(bullet.get("id", ""))],
                )
            )
        project_candidates = []
        for project_index, project in enumerate(role.get("projects", [])):
            project_bullets = []
            for bullet_index, bullet in enumerate(project.get("bullets", [])):
                project_bullets.append(
                    BulletCandidate(
                        id=f"bullet.role.{role_index}.project.{project_index}.{bullet_index}",
                        text=bullet.get("text", ""),
                        score=parse_int(bullet.get("score"), 0),
                        priority=parse_int(bullet.get("importance"), 0),
                        min_keep=bullet_index == 0,
                        drop_cost=8 if normalize_text(bullet.get("inclusion_mode", "context")) == "prefer" else 4,
                        estimated_lines=estimate_text_lines(bullet.get("text", ""), 96),
                        inclusion_mode=normalize_text(bullet.get("inclusion_mode", "context")) or "context",
                        selection_reason=["selected_project_bullet"],
                        source_ids=[str(bullet.get("id", ""))],
                    )
                )
            project_candidates.append(
                ProjectCandidate(
                    id=f"project.{role_index}.{project.get('project_id', project_index)}",
                    label=project.get("label", ""),
                    stack=project.get("stack", ""),
                    score=parse_int(project.get("score"), 0),
                    priority=parse_int(project.get("selected_point_count"), 0),
                    min_keep=project_index == 0,
                    drop_cost=12,
                    estimated_lines=1 + sum(item.estimated_lines for item in project_bullets),
                    selection_reason=["selected_project"],
                    source_ids=[str(project.get("project_id", ""))],
                    bullets=project_bullets,
                )
            )
        role_candidates.append(
            RoleCandidate(
                id=f"role.{role_index}",
                company=role.get("company", ""),
                title=role.get("title", ""),
                dates=role.get("dates", ""),
                score=sum(item.score for item in direct_bullets) + sum(project.score for project in project_candidates),
                priority=100 - role_index,
                min_keep=role_index == 0,
                drop_cost=20,
                estimated_lines=1 + len(project_candidates) + sum(item.estimated_lines for item in direct_bullets) + sum(project.estimated_lines for project in project_candidates),
                selection_reason=["selected_role"],
                source_ids=[role.get("company", ""), role.get("title", "")],
                direct_bullets=direct_bullets,
                projects=project_candidates,
            )
        )

    education_candidates = []
    for index, entry in enumerate(education_entries):
        education_candidates.append(
            EducationCandidate(
                id=f"education.{index}",
                school=entry.get("school", ""),
                detail=entry.get("detail", ""),
                year=entry.get("year", ""),
                score=20 - index,
                priority=20 - index,
                min_keep=index == 0,
                drop_cost=5,
                estimated_lines=1,
                selection_reason=["selected_education"],
                source_ids=[entry.get("school", "")],
            )
        )

    certification_candidates = []
    for index, entry in enumerate(certification_entries):
        certification_candidates.append(
            CertificationCandidate(
                id=f"certification.{index}",
                name=entry.get("name", ""),
                detail=entry.get("detail", ""),
                year=entry.get("year", ""),
                score=30 - index,
                priority=30 - index,
                min_keep=index == 0,
                drop_cost=5,
                estimated_lines=1,
                selection_reason=["selected_certification"],
                source_ids=[entry.get("name", "")],
            )
        )

    return CandidateContent(
        summary=summary_candidates,
        skill_groups=skill_group_candidates,
        roles=role_candidates,
        education=education_candidates,
        certifications=certification_candidates,
    )


def count_section_items(section_id, candidate_content):
    if section_id == "summary":
        return len(candidate_content.summary)
    if section_id == "skills":
        return len(candidate_content.skill_groups)
    if section_id == "experience":
        return len(candidate_content.roles)
    if section_id == "education":
        return len(candidate_content.education)
    if section_id == "certifications":
        return len(candidate_content.certifications)
    return 0


def estimate_section_lines(section_id, candidate_content, caps):
    if section_id == "summary":
        entries = candidate_content.summary[: caps.get("summary", len(candidate_content.summary))]
        return sum(item.estimated_lines for item in entries)
    if section_id == "skills":
        entries = candidate_content.skill_groups[: caps.get("skills", len(candidate_content.skill_groups))]
        return sum(item.estimated_lines for item in entries)
    if section_id == "experience":
        entries = candidate_content.roles[: caps.get("experience", len(candidate_content.roles))]
        return sum(item.estimated_lines for item in entries)
    if section_id == "education":
        entries = candidate_content.education[: caps.get("education", len(candidate_content.education))]
        return sum(item.estimated_lines for item in entries)
    if section_id == "certifications":
        entries = candidate_content.certifications[: caps.get("certifications", len(candidate_content.certifications))]
        return sum(item.estimated_lines for item in entries)
    return 0


def build_layout_candidates(db_data, profile_key, candidate_content, role_context, runtime_feature_map, page_format, max_pages):
    layout_rows = [dict(row) for row in (db_data.get("resume_layouts") or []) if parse_bool(row.get("approved"), True)]
    if not layout_rows:
        layout_rows = fallback_layout_rows()
    profile_rules = build_layout_profile_rule_index(db_data, profile_key)
    profile_rule_map = {row["layout_key"]: row for row in profile_rules if row.get("layout_key")}
    signal_rule_map = build_layout_signal_rule_index(db_data)
    layout_signals = build_layout_signal_map(candidate_content, runtime_feature_map)
    cert_alignment_score = parse_int(runtime_feature_map.get("cert_sensitivity"), 0)
    experience_strength = sum(item.score for item in candidate_content.roles[:2])
    skill_strength = sum(item.score for item in candidate_content.skill_groups[:3])
    education_score = parse_int(runtime_feature_map.get("education_sensitivity"), 0)
    direct_match_density = float(runtime_feature_map.get("direct_match_density") or 0.0)
    weak_context_pressure = float(runtime_feature_map.get("weak_context_pressure") or 0.0)
    role_archetype_map = build_role_archetype_score_map((runtime_feature_map or {}).get("role_archetype_scores") or [])
    allowed_lines = allowed_lines_for_plan(page_format, max_pages)

    layouts = []
    for layout_row in layout_rows:
        layout_key = layout_row.get("layout_key")
        if not layout_key:
            continue
        layout_family = normalize_text(layout_row.get("layout_family") or LAYOUT_GROUPED_STANDARD) or LAYOUT_GROUPED_STANDARD
        certification_policy = normalize_text(layout_row.get("certification_policy") or "contextual") or "contextual"
        page_count = parse_int(layout_row.get("page_count"), 1)
        if page_count > max_pages:
            continue
        if profile_rule_map and layout_key not in profile_rule_map:
            continue
        if parse_int(layout_row.get("max_roles"), 0) and len(candidate_content.roles) > parse_int(layout_row.get("max_roles"), 0):
            continue
        if parse_int(layout_row.get("allow_low_match_context"), 0) == 0 and weak_context_pressure >= 0.65 and "adjacent" not in normalize_text(layout_key):
            continue
        section_caps = build_layout_section_caps(layout_row, candidate_content)
        section_order = parse_layout_section_order(layout_row)
        section_visibility = {
            section_id: count_section_items(section_id, candidate_content) > 0
            for section_id in section_order
        }
        if certification_policy == "hidden":
            section_visibility["certifications"] = False
        estimated_lines = sum(
            estimate_section_lines(section_id, candidate_content, section_caps)
            for section_id in section_order
            if section_visibility.get(section_id, False)
        ) + 8
        estimated_pages = round(estimated_lines / allowed_lines, 2) if allowed_lines else 99.0
        profile_priority = parse_int(profile_rule_map.get(layout_key, {}).get("priority"), 0)
        score_inputs = {
            "layout_signals": dict(layout_signals),
            "profile_priority": profile_priority,
            "cert_alignment_score": cert_alignment_score,
            "experience_strength": experience_strength,
            "skill_strength": skill_strength,
            "education_score": education_score,
            "direct_match_density": direct_match_density,
            "weak_context_pressure": weak_context_pressure,
            "layout_family": layout_family,
            "certification_policy": certification_policy,
            "page_pressure_penalty": max(0, estimated_lines - allowed_lines),
        }
        role_archetype_priority = 0
        if layout_family == LAYOUT_GROUPED_TOP_CERTS or certification_policy == "top_relevant":
            role_archetype_priority += int(round(role_archetype_map.get("cyber", 0.0) * 35))
            role_archetype_priority += int(round(role_archetype_map.get("forensic", 0.0) * 40))
            role_archetype_priority += min(30, cert_alignment_score * 2)
        if layout_family == LAYOUT_GROUPED_STANDARD:
            role_archetype_priority += int(round(role_archetype_map.get("gameplay", 0.0) * 25))
            role_archetype_priority += int(round(role_archetype_map.get("liveops", 0.0) * 20))
            role_archetype_priority += int(round(role_archetype_map.get("vr", 0.0) * 18))
            role_archetype_priority += int(round(role_archetype_map.get("backend", 0.0) * 16))
        score_inputs["role_archetype_priority"] = role_archetype_priority
        score = skill_strength + experience_strength + education_score + profile_priority - score_inputs["page_pressure_penalty"]
        score += int(direct_match_density * 100)
        score -= int(weak_context_pressure * 80)
        score += role_archetype_priority
        matched_signal_rules = []
        gated = False
        for rule in signal_rule_map.get(layout_key, []):
            if not profile_signal_rule_matches(rule, layout_signals):
                continue
            weight = parse_int(row_value(rule, "weight"), 0)
            action = normalize_text(row_value(rule, "action", default="score")) or "score"
            matched_signal_rules.append(
                {
                    "signal_key": row_value(rule, "signal_key", default=""),
                    "operator": row_value(rule, "operator", default=""),
                    "weight": weight,
                    "action": action,
                }
            )
            if action == "gate":
                gated = True
            elif action == "penalty":
                score -= weight
            else:
                score += weight
        if gated:
            continue
        layouts.append(
            LayoutCandidate(
                id=layout_key,
                section_order=section_order,
                section_visibility=section_visibility,
                section_caps=section_caps,
                score_inputs=score_inputs,
                score=score,
                estimated_lines=estimated_lines,
                estimated_pages=estimated_pages,
                selection_reason=["layout_profile_rules", "layout_signal_rules"] + [f"signal:{item['signal_key']}" for item in matched_signal_rules],
            )
        )
    if not layouts:
        fallback = fallback_layout_rows()[0]
        section_caps = build_layout_section_caps(fallback, candidate_content)
        fallback_lines = sum(estimate_section_lines(section_id, candidate_content, section_caps) for section_id in LAYOUT_STABLE_SECTION_ORDER) + 8
        layouts.append(
            LayoutCandidate(
                id=fallback["layout_key"],
                section_order=list(LAYOUT_STABLE_SECTION_ORDER),
                section_visibility={section_id: True for section_id in LAYOUT_STABLE_SECTION_ORDER},
                section_caps=section_caps,
                score_inputs={"fallback": True, "layout_signals": layout_signals},
                score=0,
                estimated_lines=fallback_lines,
                estimated_pages=round(fallback_lines / allowed_lines, 2) if allowed_lines else 99.0,
                selection_reason=["layout_fallback"],
            )
        )
    return layouts


def build_fit_state(candidate_content, candidate_layouts, page_format, max_pages):
    best_layout = max(candidate_layouts, key=lambda item: (item.score, -item.estimated_lines, item.id))
    allowed_lines = allowed_lines_for_plan(page_format, max_pages)
    sufficiency = ContentSufficiency(
        summary=sufficiency_bucket(len(candidate_content.summary)),
        skills=sufficiency_bucket(len(candidate_content.skill_groups)),
        experience=sufficiency_bucket(len(candidate_content.roles)),
        certifications=sufficiency_bucket(len(candidate_content.certifications)),
    )
    return FitState(
        page_pressure=pressure_bucket(best_layout.estimated_lines, allowed_lines),
        estimated_lines=best_layout.estimated_lines,
        estimated_pages=best_layout.estimated_pages,
        content_sufficiency=sufficiency,
        rebalance_actions=[],
    )


def rebalance_plan_to_fit(candidate_content, candidate_layouts, fit_state, page_format, max_pages):
    allowed_lines = allowed_lines_for_plan(page_format, max_pages)
    dropped_items = []
    actions = []
    selected_layout = max(candidate_layouts, key=lambda item: (item.score, -item.estimated_lines, item.id))
    standard_layout_id = next((item.id for item in candidate_layouts if "standard-1p" in item.id), selected_layout.id)

    if fit_state.page_pressure in {"high", "critical"}:
        if "adjacent" not in normalize_text(selected_layout.id) and fit_state.content_sufficiency.certifications in {"weak", "missing"} and any("standard-1p" in item.id for item in candidate_layouts):
            actions.append(
                RebalanceAction(
                    action="switch_layout",
                    target_id=standard_layout_id,
                    reason="content_insufficient",
                    details={"from": selected_layout.id, "to": standard_layout_id},
                )
            )
            selected_layout = next(item for item in candidate_layouts if item.id == standard_layout_id)

        certification_policy = normalize_text(selected_layout.score_inputs.get("certification_policy")) or "contextual"
        if certification_policy != "top_relevant" and fit_state.content_sufficiency.certifications != "strong" and candidate_content.certifications:
            dropped_items.extend(
                DroppedItem(
                    id=item.id,
                    kind="certification",
                    reason="page_pressure",
                    score=item.score,
                    details={"layout": selected_layout.id},
                )
                for item in candidate_content.certifications
            )
            actions.append(
                RebalanceAction(
                    action="hide_certifications",
                    target_id="section.certifications",
                    reason="page_pressure",
                    details={"count": len(candidate_content.certifications)},
                )
            )

        for role in candidate_content.roles[1:]:
            if role.min_keep:
                continue
            if selected_layout.estimated_lines <= allowed_lines:
                break
            dropped_items.append(
                DroppedItem(
                    id=role.id,
                    kind="role",
                    reason="higher_value_content_preserved",
                    score=role.score,
                    details={"company": role.company, "title": role.title},
                )
            )
            actions.append(
                RebalanceAction(
                    action="drop_role",
                    target_id=role.id,
                    reason="page_pressure",
                    details={"score": role.score},
                )
            )
            selected_layout.estimated_lines = max(0, selected_layout.estimated_lines - role.estimated_lines)

    fit_state.page_pressure = pressure_bucket(selected_layout.estimated_lines, allowed_lines)
    fit_state.estimated_lines = selected_layout.estimated_lines
    fit_state.estimated_pages = round(selected_layout.estimated_lines / allowed_lines, 2) if allowed_lines else 99.0
    fit_state.rebalance_actions = actions
    return selected_layout, dropped_items, fit_state


def finalize_resume_plan(archetype_name, profile_key, page_format, max_pages, inputs, candidate_content, candidate_layouts, fit_state, selected_layout, dropped_items):
    final_sections = []
    for order, section_id in enumerate(selected_layout.section_order, start=1):
        visible = selected_layout.section_visibility.get(section_id, False)
        if section_id == "certifications" and any(item.kind == "certification" for item in dropped_items):
            visible = False
        final_sections.append(
            PlannedSection(
                section_id=section_id,
                visible=visible,
                order=order,
                layout_reason=["layout_order_applied"],
                estimated_lines=estimate_section_lines(section_id, candidate_content, selected_layout.section_caps) if visible else 0,
                item_ids=[f"{section_id}.{index}" for index in range(count_section_items(section_id, candidate_content))] if visible else [],
            )
        )

    validation = PlanValidation(
        passed=fit_state.page_pressure != "critical",
        violations=[] if fit_state.page_pressure != "critical" else ["estimated_layout_exceeds_page_budget"],
    )

    return ResumePlan(
        version="1.0",
        archetype=archetype_name,
        profile_key=profile_key,
        format=page_format,
        max_pages=max_pages,
        inputs=inputs,
        candidate_content=candidate_content,
        candidate_layouts=candidate_layouts,
        fit_state=fit_state,
        final_layout=selected_layout.id,
        final_sections=final_sections,
        dropped_items=dropped_items,
        validation=validation,
    )


def explain_profile_selection(resume_profile):
    candidates = list((resume_profile or {}).get("profile_candidates") or [])
    selected_key = normalize_resume_profile_key((resume_profile or {}).get("profile_key", ""))
    selected = next((item for item in candidates if normalize_resume_profile_key(item.get("profile_key")) == selected_key), None)
    return {
        "selected_profile_key": selected_key,
        "fit_tier": (resume_profile or {}).get("fit_tier"),
        "selection_reason": (resume_profile or {}).get("selection_reason"),
        "winning_candidate": selected,
        "all_candidates": candidates,
        "derived_signals": (resume_profile or {}).get("profile_signals", {}),
    }


def explain_layout_selection(resume_plan):
    if not resume_plan:
        return {}
    plan = asdict(resume_plan) if hasattr(resume_plan, "__dataclass_fields__") else dict(resume_plan)
    final_layout = plan.get("final_layout")
    candidates = plan.get("candidate_layouts", [])
    winning = next((item for item in candidates if item.get("id") == final_layout), None)
    return {
        "selected_layout": final_layout,
        "winning_candidate": winning,
        "candidate_layouts": candidates,
        "rebalance_actions": plan.get("fit_state", {}).get("rebalance_actions", []),
        "validation": plan.get("validation", {}),
    }


def explain_experience_selection(roles, resume_plan, experience_selection_budget=None):
    plan = asdict(resume_plan) if hasattr(resume_plan, "__dataclass_fields__") else (resume_plan or {})
    dropped_items = plan.get("dropped_items", []) if isinstance(plan, dict) else []
    annotated_drops = []
    for item in dropped_items:
        reason = item.get("reason")
        details = dict(item.get("details") or {})
        if item.get("kind") == "role":
            details["explanation"] = "Dropped because higher-scoring roles consumed the available page/role budget."
        elif item.get("kind") == "certification":
            details["explanation"] = "Dropped because layout/page pressure preserved stronger experience content."
        annotated_drops.append({**item, "details": details, "human_reason": reason})
    kept_role_lookup = {
        (role.get("company"), role.get("title")): role
        for role in roles
    }
    candidate_roles = ((plan.get("candidate_content") or {}).get("roles") or []) if isinstance(plan, dict) else []
    role_decisions = []
    for candidate_role in candidate_roles:
        role_key = (candidate_role.get("company"), candidate_role.get("title"))
        kept_role = kept_role_lookup.get(role_key)
        kept = kept_role is not None
        kept_project_ids = {project.get("project_id") for project in kept_role.get("projects", [])} if kept_role else set()
        kept_direct_ids = {bullet.get("id") for bullet in kept_role.get("direct_bullets", [])} if kept_role else set()

        project_decisions = []
        for candidate_project in candidate_role.get("projects", []):
            kept_project = kept and candidate_project.get("source_ids", [None])[0] and str(candidate_project.get("source_ids", [None])[0]) == str(candidate_project.get("source_ids", [None])[0])
            project_id = None
            source_ids = candidate_project.get("source_ids") or []
            if source_ids:
                try:
                    project_id = int(source_ids[0])
                except Exception:
                    project_id = source_ids[0]
            included = project_id in kept_project_ids if project_id is not None else False
            bullet_decisions = []
            kept_project_bullets = []
            if kept_role:
                for project in kept_role.get("projects", []):
                    if project.get("project_id") == project_id:
                        kept_project_bullets = project.get("bullets", [])
                        break
            kept_project_bullet_ids = {bullet.get("id") for bullet in kept_project_bullets}
            for candidate_bullet in candidate_project.get("bullets", []):
                bullet_id = None
                source_ids = candidate_bullet.get("source_ids") or []
                if source_ids:
                    bullet_id = source_ids[0]
                bullet_included = bullet_id in kept_project_bullet_ids
                bullet_decisions.append(
                    {
                        "id": bullet_id,
                        "text": candidate_bullet.get("text"),
                        "score": candidate_bullet.get("score"),
                        "included": bullet_included,
                        "reason": (
                            "Kept because it fit within the project bullet budget."
                            if bullet_included
                            else "Dropped because stronger bullets consumed the project bullet budget."
                        ),
                    }
                )
            project_decisions.append(
                {
                    "project_id": project_id,
                    "label": candidate_project.get("label"),
                    "score": candidate_project.get("score"),
                    "included": included,
                    "reason": (
                        "Kept because the project survived role/project ranking."
                        if included
                        else "Dropped because higher-value projects consumed the project budget."
                    ),
                    "bullets": bullet_decisions,
                }
            )

        direct_bullet_decisions = []
        for candidate_bullet in candidate_role.get("direct_bullets", []):
            bullet_id = None
            source_ids = candidate_bullet.get("source_ids") or []
            if source_ids:
                bullet_id = source_ids[0]
            included = bullet_id in kept_direct_ids
            direct_bullet_decisions.append(
                {
                    "id": bullet_id,
                    "text": candidate_bullet.get("text"),
                    "score": candidate_bullet.get("score"),
                    "included": included,
                    "reason": (
                        "Kept because it survived direct bullet ranking."
                        if included
                        else "Dropped because stronger direct/project bullets consumed the bullet budget."
                    ),
                }
            )

        role_decisions.append(
            {
                "company": candidate_role.get("company"),
                "title": candidate_role.get("title"),
                "score": candidate_role.get("score"),
                "included": kept,
                "reason": (
                    "Kept because the role survived ranked experience selection budget."
                    if kept
                    else "Dropped because higher-scoring roles consumed the role/page budget."
                ),
                "direct_bullets": direct_bullet_decisions,
                "projects": project_decisions,
            }
        )
    return {
        "selection_budget": experience_selection_budget or {},
        "role_decisions": role_decisions,
        "dropped_items": annotated_drops,
    }


def build_resume_plan(db_data, archetype_name, resume_profile, args, role_context, runtime_feature_map, skill_groups, roles, education_entries, certification_entries):
    max_pages = parse_int(getattr(args, "max_pages", 1), 1)
    profile_key = normalize_resume_profile_key(resume_profile.get("profile_key", archetype_name))
    candidate_content = build_candidate_content(
        resume_profile.get("summary", ""),
        skill_groups,
        roles,
        education_entries,
        certification_entries,
    )
    candidate_layouts = build_layout_candidates(db_data, profile_key, candidate_content, role_context, runtime_feature_map, args.format, max_pages)
    fit_state = build_fit_state(candidate_content, candidate_layouts, args.format, max_pages)
    selected_layout, dropped_items, fit_state = rebalance_plan_to_fit(
        candidate_content,
        candidate_layouts,
        fit_state,
        args.format,
        max_pages,
    )
    return finalize_resume_plan(
        archetype_name,
        profile_key,
        args.format,
        max_pages,
        {
            "keywords": [item.strip() for item in (args.keywords or "").split(",") if item.strip()],
            "role_context": role_context,
            "runtime_feature_map": runtime_feature_map,
        },
        candidate_content,
        candidate_layouts,
        fit_state,
        selected_layout,
        dropped_items,
    )


def write_resume_metadata(output_dir, archetype_name, keywords, resume_profile, db_data, roles, skill_groups, education_entries, certification_entries, role_context=None, runtime_feature_map=None, experience_selection_budget=None, resume_plan=None):
    explainability = {
        "derived_features": runtime_feature_map or {},
        "profile_selection": explain_profile_selection(resume_profile),
        "layout_selection": explain_layout_selection(resume_plan),
        "education_selection": explain_education_selection(db_data, education_entries, runtime_feature_map=runtime_feature_map),
        "certification_selection": explain_certification_selection(db_data, certification_entries, runtime_feature_map=runtime_feature_map),
        "experience_selection": explain_experience_selection(roles, resume_plan, experience_selection_budget=experience_selection_budget),
    }
    payload = {
        "archetype": archetype_name,
        "keywords": keywords,
        "resume_profile": resume_profile,
        "role_archetype_scores": (runtime_feature_map or {}).get("role_archetype_scores") or [],
        "db_rule_context": {
            "resume_profiles": len(db_data.get("resume_profiles", [])),
            "resume_profile_signal_rules": len(db_data.get("resume_profile_signal_rules", [])),
            "resume_point_variants": len(db_data.get("resume_point_variants", [])),
            "skill_resume_rules": len(db_data.get("skill_resume_rules", [])),
            "resume_group_rules": len(db_data.get("resume_group_rules", [])),
        },
        "role_context": role_context or {},
        "runtime_feature_map": runtime_feature_map or {},
        "experience_selection_budget": experience_selection_budget or {},
        "explainability": explainability,
        "skill_groups": skill_groups,
        "roles": roles,
        "education": education_entries,
        "certifications": certification_entries,
        "resume_plan": resume_plan,
    }
    path = output_dir / "resume-metadata.json"
    path.write_text(json.dumps(make_json_safe(payload), indent=2), encoding="utf-8")
    return path


def render_resume(args):
    db_data = load_db()
    try:
        profile = load_yaml(profile_path())
    except SystemExit:
        profile = {"candidate": {}}
    requested_archetype = normalize_requested_archetype(args.archetype)
    valid_archetypes = set((db_data.get("archetype_resume_configs") or {}).keys())
    if requested_archetype not in valid_archetypes:
        fail(f"Invalid --archetype '{args.archetype}'. Valid: {', '.join(sorted(valid_archetypes))}")
    archetype = build_archetype_config(requested_archetype, db_data=db_data)

    candidate = dict(profile.get("candidate") or {})
    candidate.setdefault("full_name", db_data["profile"].get("full_name", ""))
    candidate.setdefault("display_name", db_data["profile"].get("display_name", candidate.get("full_name", "")))
    candidate.setdefault("email", db_data["profile"].get("email", ""))
    candidate.setdefault("phone", db_data["profile"].get("phone", ""))
    candidate.setdefault("location", db_data["profile"].get("location", ""))

    keywords = [item.strip() for item in (args.keywords or "").split(",") if item.strip()]
    role_context = resolve_role_context(db_data, args)
    resume_profile = resolve_resume_profile(db_data, requested_archetype, archetype, keywords=keywords, role_context=role_context)
    roles, experience_selection_budget = build_selected_roles(
        db_data,
        requested_archetype,
        archetype,
        keywords,
        parse_int(getattr(args, "max_pages", 1), 1),
        resume_profile=resume_profile,
    )
    education_entries = []
    certification_entries = []
    runtime_feature_map = build_runtime_feature_map(
        db_data,
        role_context,
        keywords,
        roles,
        education_entries,
        certification_entries,
    )
    education_entries = build_education_entries(db_data, archetype, keywords, runtime_feature_map=runtime_feature_map)
    certification_entries = build_certification_entries(db_data, archetype, keywords, runtime_feature_map=runtime_feature_map)
    runtime_feature_map = build_runtime_feature_map(
        db_data,
        role_context,
        keywords,
        roles,
        education_entries,
        certification_entries,
    )
    skill_groups = build_skill_groups(
        db_data,
        db_data["skills"],
        requested_archetype,
        archetype,
        keywords,
        role_context=role_context,
        runtime_feature_map=runtime_feature_map,
        max_pages=parse_int(getattr(args, "max_pages", 1), 1),
    )
    resume_plan = build_resume_plan(
        db_data,
        requested_archetype,
        resume_profile,
        args,
        role_context,
        runtime_feature_map,
        skill_groups,
        roles,
        education_entries,
        certification_entries,
    )

    template_key = archetype.get("_meta", {}).get("template_key") or requested_archetype
    template_path = CONFIG_DIR / "resume-templates" / f"{template_key}-template.docx"
    if not template_path.exists():
        fallback = CONFIG_DIR / "resume-templates" / "general-template.docx"
        if template_key in {"frontend", "gameserver"} and fallback.exists():
            template_path = fallback
        else:
            fail(f"Missing template: {template_path}")

    document = Document(str(template_path))
    apply_page_format(document, args.format)
    replace_placeholders(
        document,
        {
            "{{NAME}}": candidate.get("full_name", ""),
            "{{SUBTITLE}}": resume_profile["subtitle"],
            "{{EMAIL}}": candidate.get("email", ""),
            "{{PHONE}}": candidate.get("phone", ""),
            "{{LOCATION}}": candidate.get("location", ""),
            "{{SUMMARY}}": resume_profile["summary"],
        },
    )

    donors = find_donors(document)
    build_skills_section(document, donors, skill_groups)
    build_experience_section(document, donors, roles, is_games=bool(runtime_feature_map.get("is_games")))
    build_education_section(document, donors, education_entries)
    append_certifications_section(document, donors, certification_entries)
    apply_section_plan_to_document(document, resume_plan)
    normalize_section_columns(document)

    output_dir = Path(args.out) if args.out else OUTPUT_DIR / f"{requested_archetype}-resume"
    if not output_dir.is_absolute():
        output_dir = (REPO_ROOT / output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    name_prefix = (candidate.get("full_name") or "").replace(" ", "") or "Candidate"
    docx_path = output_dir / f"{name_prefix}-resume.docx"
    pdf_path = output_dir / f"{name_prefix}-resume.pdf"
    write_resume_metadata(
        output_dir,
        requested_archetype,
        keywords,
        resume_profile,
        db_data,
        roles,
        skill_groups,
        education_entries,
        certification_entries,
        role_context=role_context,
        runtime_feature_map=runtime_feature_map,
        experience_selection_budget=experience_selection_budget,
        resume_plan=resume_plan,
    )
    document.save(str(docx_path))

    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
    except Exception:
        pass

    print(str(pdf_path if pdf_path.exists() else docx_path))
    return {"docx": docx_path, "pdf": pdf_path if pdf_path.exists() else None}


def parse_args(argv):
    parser = argparse.ArgumentParser()
    parser.add_argument("--archetype", required=True)
    parser.add_argument("--keywords", default="")
    parser.add_argument("--format", default="letter", choices=["letter", "a4"])
    parser.add_argument("--max-pages", default="1", choices=["1", "2"])
    parser.add_argument("--out", default="")
    parser.add_argument("--job-role", default="")
    parser.add_argument("--job-company", default="")
    parser.add_argument("--job-notes", default="")
    parser.add_argument("--job-location", default="")
    parser.add_argument("--job-work-model", default="")
    parser.add_argument("--job-compensation", default="")
    return parser.parse_args(argv)


def main():
    args = parse_args(sys.argv[1:])
    render_resume(args)


if __name__ == "__main__":
    main()
